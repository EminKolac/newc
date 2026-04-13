#!/usr/bin/env python3
"""
BBB DOCX Rapor Üretici v3.0 — Kurum seviyesi profesyonel equity analiz raporları.

Özellikler:
  - Dual-column layout (kenarlıksız tablo grid)
  - AK Yatırım tarzı Sayfa 1 (sol tez + sağ summary box)
  - 2'li grafik grid (yan yana chart yerleşimi)
  - Finansal tablo grid (gelir tablosu + nakit akış yan yana)
  - Header/footer (şirket adı + sayfa numarası)
  - Gerçek İçindekiler (Word TOC field)
  - Grafik boyut kontrolü (TAM/YARIM/GRID)
  - BBB renk paleti (#f7931a turuncu) + Calibri
  - 300 DPI grafik embed
  - Yasal uyarı

Kullanım:
    # AK Yatırım tarzı initiation şablonu
    python rapor-uret.py --sablon initiation --ticker TBORG --sirket "Türk Tuborg" --cikti rapor.docx

    # Markdown → DOCX (profesyonel layout)
    python rapor-uret.py --markdown analysis.md --ticker TBORG --cikti rapor.docx --grafikler charts/

    # Earnings update
    python rapor-uret.py --sablon earnings --ticker TBORG --cikti Q4_update.docx

Detaylı DOCX standartları: references/c2-tam-kapsama/profesyonel-cikti-rehberi.md
"""

import argparse
import os
import sys
import re
from pathlib import Path
from typing import List, Optional, Dict, Any

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml, OxmlElement
except ImportError:
    print("python-docx gerekli. Kur: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# ═══════════════════════════════════════════════════════════
# BBB STİL SABİTLERİ
# ═══════════════════════════════════════════════════════════

class BBB:
    """BBB renk paleti ve stil sabitleri — Sıcak Premium tasarım dili."""
    # Renkler
    BIRINCIL = RGBColor(0xf7, 0x93, 0x1a)      # BBB turuncu
    IKINCIL = RGBColor(0x4d, 0x4d, 0x4d)        # Koyu gri
    ACIK = RGBColor(0xBD, 0xD7, 0xEE)           # Açık mavi
    BEYAZ = RGBColor(0xFF, 0xFF, 0xFF)
    METIN = RGBColor(0x33, 0x33, 0x33)
    GRI = RGBColor(0x80, 0x80, 0x80)
    ACIK_GRI = RGBColor(0xED, 0xE7, 0xDA)       # Sıcak bej-gri (bej sayfa uyumlu)
    SIYAH = RGBColor(0x00, 0x00, 0x00)
    OLUMLU = RGBColor(0x54, 0x82, 0x35)         # Yeşil
    OLUMSUZ = RGBColor(0xC0, 0x00, 0x00)        # Kırmızı

    # Sıcak Premium palet eklentileri
    SAYFA_BG = RGBColor(0xFF, 0xFF, 0xFF)       # v4.13: Beyaz sayfa arka planı — grafik PNG'ler ile uyumlu
    SICAK_GRI = RGBColor(0x8C, 0x7E, 0x6E)      # Warm gray — kaynak notları, ikincil metin
    KOYU_TURUNCU = RGBColor(0xD4, 0x7A, 0x15)   # Koyu turuncu — hover/vurgu alternatif

    # Hex karşılıkları (tablo shading için)
    BIRINCIL_HEX = "f7931a"
    IKINCIL_HEX = "4d4d4d"
    ACIK_HEX = "FFF3E0"
    ACIK_GRI_HEX = "EDE7DA"        # Sıcak bej-gri (grafik kutucuk header, zebra-stripe)
    SAYFA_BG_HEX = "FFFFFF"        # v4.13: Beyaz sayfa arka plan
    KREM_HEX = "FFF8F0"            # Açık krem — tablo alternatif satır

    # Tavsiye renk etiketleri (kapak)
    TAVSIYE_RENK = {
        'AL': 'AL_YESIL',
        'EKLE': 'AL_YESIL',
        'TUT': 'TUT_TURUNCU',
        'SAT': 'SAT_KIRMIZI',
        'AZALT': 'SAT_KIRMIZI',
    }
    AL_HEX = "548235"
    TUT_HEX = "f7931a"
    SAT_HEX = "C00000"

    # Font — top-tier yatırım bankası standardı (JPM/GS referans: Arial)
    FONT = 'Arial'
    BASLIK_PT = 14              # H1 — güçlü tipografi hiyerarşisi (gövdeden +5pt fark)
    ALT_BASLIK_PT = 11          # H2 — gövdeden +2pt, H1'den net ayrışma
    METIN_PT = 9                # Gövde — 9pt = sektör standardı
    TABLO_PT = 8
    KAYNAK_PT = 7.5
    KUCUK_PT = 7

    # Sayfa — agresif margin = maksimum içerik alanı (top-tier standart: AK Yatırım referans)
    # v4.8: 2.0/1.8 → 1.5/1.2 — referans kurumlarla uyumlu, sağ ekstra sıkı
    SOL_KENAR_CM = 1.5
    SAG_KENAR_CM = 1.2
    UST_KENAR_CM = 1.2
    ALT_KENAR_CM = 1.2
    KENAR_CM = 1.5  # geriye uyumluluk
    SAYFA_GENISLIK_CM = 18.3  # A4 (21cm) - sol(1.5) - sağ(1.2) = 18.3cm

    # Grafik boyutları — kompakt, sayfa yoğunluğu odaklı (v4.8: yeni margin'lara uyumlu)
    GRAFIK_TAM = Cm(17.5)       # Tam genişlik — sayfa genişliğine yakın
    GRAFIK_ORTA = Cm(13.0)      # Tekil grafik — sayfanın %70'i, nefes alanı bırakır
    GRAFIK_YARIM = Cm(8.5)      # 2'li yan yana grid
    GRAFIK_GRID = Cm(8.0)       # NxM grid hücresi
    GRAFIK_KUCUK = Cm(5.0)      # Küçük inline grafik
    GRAFIK_PASTA = Cm(5.5)      # v4.11: Pasta/donut grafik — dikey alan sınırlı (S.7, S.9 fix)

    # Hyperlink renk
    LINK = RGBColor(0x05, 0x63, 0xC1)  # Word mavi


# ─────────────────────────────────────────────────────────────
# BBB_RENKLER — hex string paleti (T5 primitifleri için)
# ─────────────────────────────────────────────────────────────
BBB_RENKLER = {
    'birincil': 'f7931a',   # BBB turuncu (# olmadan — _hucre_shading uyumlu)
    'ikincil': '4d4d4d',
    'acik': 'FFF3E0',
    'acik_gri': 'EDE7DA',   # Sıcak bej-gri (bej sayfa uyumlu)
    'sayfa_bg': 'F5F0E8',   # Sayfa arka plan bej
    'krem': 'FFF8F0',       # Açık krem — tablo alternatif satır
    'olumlu': '548235',
    'olumsuz': 'C00000',
}


# ═══════════════════════════════════════════════════════════
# SIRASAL NUMARALAMA SİSTEMİ
# ═══════════════════════════════════════════════════════════

class _RaporSayaci:
    """Grafik ve Tablo sıralı numaralama. Rapor başında sıfırlanır."""
    def __init__(self):
        self.grafik = 0
        self.tablo = 0

    def sonraki_grafik(self):
        self.grafik += 1
        return self.grafik

    def sonraki_tablo(self):
        self.tablo += 1
        return self.tablo

    def sifirla(self):
        self.grafik = 0
        self.tablo = 0

_sayac = _RaporSayaci()


def sayac_sifirla():
    """Rapor başında grafik/tablo sayacını sıfırla."""
    _sayac.sifirla()


# ═══════════════════════════════════════════════════════════
# TEMEL YARDIMCI FONKSİYONLAR
# ═══════════════════════════════════════════════════════════

def add_hyperlink(paragraph, display_text, url, font_size=None, bold=False):
    """Word-native tıklanabilir hyperlink ekle.

    Kullanım:
        p = doc.add_paragraph()
        add_hyperlink(p, 'KAP Bildirim', 'https://kap.org.tr/tr/Bildirim/12345')
    """
    part = paragraph.part
    r_id = part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Stil: mavi + altı çizili
    c_elem = OxmlElement('w:color')
    c_elem.set(qn('w:val'), '0563C1')
    rPr.append(c_elem)

    u_elem = OxmlElement('w:u')
    u_elem.set(qn('w:val'), 'single')
    rPr.append(u_elem)

    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))  # half-points
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(szCs)

    if bold:
        b_elem = OxmlElement('w:b')
        rPr.append(b_elem)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), BBB.FONT)
    rFonts.set(qn('w:hAnsi'), BBB.FONT)
    rPr.append(rFonts)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = display_text
    text_elem.set(qn('xml:space'), 'preserve')
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
    return hyperlink


def kaynaklar_sayfasi(doc, kaynaklar):
    """Son sayfa: Kaynaklar & Referanslar.

    kaynaklar: dict of {kategori: [(display, url, tarih), ...]}

    Örnek:
        kaynaklar_sayfasi(doc, {
            'Finansal Veriler': [
                ('KAP Q4 2025 Finansal Tablo', 'https://kap.org.tr/...', '15 Subat 2026'),
                ('BBB Finans API ciktisi', None, '18 Mart 2026'),
            ],
            'Degerleme Referanslari': [
                ('Damodaran Country Risk Premium', 'https://pages.stern.nyu.edu/~adamodar/...', 'Mart 2026'),
            ],
        })
    """
    # v4.11: Kaynaklar her zaman yeni sayfada başlamalı
    _baslik(doc, 'Kaynaklar & Referanslar', level=1, sayfa_sonu=True)

    for kategori, items in kaynaklar.items():
        _paragraf(doc, kategori, size=BBB.METIN_PT, bold=True, color=BBB.BIRINCIL,
                  space_after=4, space_before=8)
        for item in items:
            display = item[0]
            url = item[1] if len(item) > 1 else None
            tarih = item[2] if len(item) > 2 else None

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5)

            # Bullet
            run_bullet = p.add_run('• ')
            _font(run_bullet, size=BBB.KAYNAK_PT, color=BBB.METIN)

            if url:
                add_hyperlink(p, display, url, font_size=BBB.KAYNAK_PT)
            else:
                run_text = p.add_run(display)
                _font(run_text, size=BBB.KAYNAK_PT, color=BBB.METIN)

            if tarih:
                run_date = p.add_run(f'  (Erisim: {tarih})')
                _font(run_date, size=BBB.KAYNAK_PT, italic=True, color=BBB.GRI)


def _font(run, size=BBB.METIN_PT, bold=False, italic=False, color=None, name=BBB.FONT):
    """Run font stilini ayarla. color: RGBColor objesi veya '#RRGGBB' / 'RRGGBB' hex string."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        if isinstance(color, str):
            hex_str = color.lstrip('#')
            color = RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))
        run.font.color.rgb = color
    # CJK uyumluluğu
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def _paragraf(doc_or_cell, text, size=BBB.METIN_PT, bold=False, italic=False,
              color=None, alignment=None, space_after=None, space_before=None,
              line_spacing=None, inline_markdown=False):
    """Stillenmiş paragraf ekle. doc veya table cell kabul eder.

    alignment varsayılanı: JUSTIFY (profesyonel rapor standardı).
    line_spacing: Pt cinsinden satır aralığı (varsayılan: Pt(12) ≈ 1.2× for 10pt).
    inline_markdown: True ise **bold** ve *italic* markdown parse edilir (v4.14).
                     Bold lead-in için: '**Bilanço güçlüdür:** Nakit 2,34 mrd TL...'
    """
    if hasattr(doc_or_cell, 'add_paragraph'):
        if hasattr(doc_or_cell, 'vertical_alignment'):
            # Table cell — ilk çağrıda varsayılan boş paragrafı yeniden kullan,
            # sonraki çağrılarda yeni paragraf ekle (içerik kaybını önle)
            first_p = doc_or_cell.paragraphs[0]
            if first_p.text.strip() == '' and not first_p.runs:
                p = first_p
            else:
                p = doc_or_cell.add_paragraph()
        else:
            p = doc_or_cell.add_paragraph()
    else:
        # Table cell fallback
        if doc_or_cell.paragraphs:
            first_p = doc_or_cell.paragraphs[0]
            if first_p.text.strip() == '' and not first_p.runs:
                p = first_p
            else:
                p = doc_or_cell.add_paragraph()
        else:
            p = doc_or_cell.add_paragraph()
    # Varsayılan JUSTIFY — explicit None ile override edilebilir
    if alignment is not None:
        p.alignment = alignment
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = Pt(space_after)
    else:
        pf.space_after = Pt(5)  # v4.11: 2→5pt — paragraflar arası görünür boşluk (kapak stili tutarlılık)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    # Satır aralığı — AT_LEAST kuralı: metin için minimum, grafik için genişleyebilir
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    else:
        pf.line_spacing = Pt(10.5)  # ~1.17× for 9pt — top-tier sıkı yoğunluk
    pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    if inline_markdown and ('**' in text or '*' in text):
        # v4.14: Inline markdown parsing — bold lead-in desteği
        import re as _re
        _pattern = _re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))')
        for _m in _pattern.finditer(text):
            if _m.group(2):  # ***bold italic***
                _run = p.add_run(_m.group(2))
                _font(_run, size=size, bold=True, italic=True, color=color or BBB.METIN)
            elif _m.group(3):  # **bold**
                _run = p.add_run(_m.group(3))
                _font(_run, size=size, bold=True, color=color or BBB.METIN)
            elif _m.group(4):  # *italic*
                _run = p.add_run(_m.group(4))
                _font(_run, size=size, italic=True, color=color or BBB.METIN)
            elif _m.group(5):  # normal text
                _run = p.add_run(_m.group(5))
                _font(_run, size=size, bold=bold, italic=italic, color=color or BBB.METIN)
    else:
        run = p.add_run(text)
        _font(run, size=size, bold=bold, italic=italic, color=color or BBB.METIN)
    return p


def _baslik(doc, text, level=1, sayfa_sonu=None):
    """BBB stilinde heading — Sıcak Premium görsel hiyerarşi.

    H1: 14pt bold turuncu + üst/alt çift çizgi bandı (MS Blue Line tarzı)
    H2: 11pt bold turuncu + sol bar (4pt tema rengi)
    H3: 9pt bold + kompakt

    Args:
        sayfa_sonu: None → H1 için False (default — akış korunur), True/False → explicit override.
                    Sadece büyük ana bölüm geçişlerinde True kullanılmalıdır.
                    v4.7: Varsayılan True→False değiştirildi (sayfa doluluk optimizasyonu).
    """
    # level → font boyutu eşlemesi
    size_map = {1: BBB.BASLIK_PT, 2: BBB.ALT_BASLIK_PT, 3: BBB.METIN_PT}
    size = size_map.get(level, BBB.METIN_PT)

    p = doc.add_paragraph()
    # v4.13: Spacing artırıldı — bloklar arası okunabilir boşluk
    if level == 1:
        p.paragraph_format.space_before = Pt(18)   # v4.13: 10→18pt — H1 öncesi belirgin ayrım
        p.paragraph_format.space_after = Pt(6)     # v4.13: 3→6pt — H1 sonrası içerikle nefes alanı
        do_break = sayfa_sonu if sayfa_sonu is not None else False
        p.paragraph_format.page_break_before = do_break
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)   # v4.13: 6→12pt — H2 öncesi belirgin ayrım
        p.paragraph_format.space_after = Pt(4)     # v4.13: 2→4pt
    else:
        p.paragraph_format.space_before = Pt(6)    # v4.13: 4→6pt
        p.paragraph_format.space_after = Pt(2)     # v4.13: 1→2pt

    # ── Sayfa kırma koruması: başlık asla sonraki içerikten kopmamalı ──
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.widow_control = True

    # Heading stilini referans için uygula (numaralandırma/TOC için)
    try:
        style_name = f'Heading {level}'
        p.style = doc.styles[style_name]
    except KeyError:
        pass  # Stil yoksa devam et

    run = p.add_run(text)
    # rPr üzerinden doğrudan renk yaz — Word style override'ını aşar
    rPr = run._element.get_or_add_rPr()
    color_elem = rPr.find(qn('w:color'))
    if color_elem is None:
        color_elem = OxmlElement('w:color')
        rPr.append(color_elem)
    color_elem.set(qn('w:val'), BBB.BIRINCIL_HEX)

    run.font.name = BBB.FONT
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.italic = False

    # ── Görsel Hierarchy: Border elementleri ──
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    if level == 1:
        # H1: Üst + alt çizgi bandı — turuncu çerçeve (Sıcak Premium imza)
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '8')    # 8 = 1pt — ince üst çizgi
        top.set(qn('w:space'), '3')
        top.set(qn('w:color'), BBB.BIRINCIL_HEX)
        pBdr.append(top)

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # 12 = 1.5pt — kalın alt çizgi
        bottom.set(qn('w:space'), '3')
        bottom.set(qn('w:color'), BBB.BIRINCIL_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        # H2: Sol bar — 3pt tema rengi
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')  # 24 = 3pt
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), BBB.BIRINCIL_HEX)
        pBdr.append(left)
        pPr.append(pBdr)
        p.paragraph_format.left_indent = Cm(0.3)

    return p


def _kaynak(doc_or_cell, text='Kaynak: KAP, BBB tahminleri'):
    """Kaynak notu (italic, gri, 8pt, sol hizalı — tablo sol kenarıyla hizalı)."""
    return _paragraf(doc_or_cell, text, size=BBB.KAYNAK_PT, italic=True, color=BBB.GRI,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=1, space_after=10)  # v4.14: 6→10pt — tablo sonrası yeterli boşluk


def _baslik_bar(doc_or_cell, text, genislik_cm=None):
    """Standart grafik/tablo başlık barı — bej-gri arka plan + koyu metin.

    v4.11: Tüm grafik ve tablo başlıkları bu fonksiyonla oluşturulur.
    grafik_kutucuk() stilini tüm rapora taşır — tutarlı görsel dil.
    v4.14: genislik_cm parametresi — tablo ile simetrik genişlik sağlar.

    doc_or_cell: Document veya table cell — her ikisinde de çalışır.
    genislik_cm: Belirtilirse başlık barı bu genişlikte çizilir (tablo ile simetri).
                 None ise sayfa genişliğinde (eski davranış).
    """
    # 1×1 mini tablo ile arka plan rengini güvenilir şekilde uygula
    if hasattr(doc_or_cell, 'add_table'):
        tbl = doc_or_cell.add_table(rows=1, cols=1)
    else:
        # Cell içinde tablo ekle
        tbl = doc_or_cell.add_table(rows=1, cols=1)

    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    header_cell = tbl.cell(0, 0)
    _hucre_shading(header_cell, BBB.ACIK_GRI_HEX)
    header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = _paragraf(header_cell, text, size=BBB.TABLO_PT, bold=True, color=BBB.IKINCIL,
                  space_after=4, space_before=2, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  line_spacing=Pt(10))  # v4.14: space_after 2→4pt — tablo ile başlık barı arası boşluk

    # Kenarlıkları kaldır — sadece arka plan rengi kalmalı
    _tablo_kenarsiz(tbl)

    # v4.14: Genişlik kontrolü — tablo ile simetri
    tbl.autofit = False
    tbl.allow_autofit = False
    if genislik_cm:
        _tbl_pr = tbl._tbl.tblPr
        for _old in _tbl_pr.findall(qn('w:tblW')):
            _tbl_pr.remove(_old)
        _w = OxmlElement('w:tblW')
        _w.set(qn('w:w'), str(int(genislik_cm * 567)))
        _w.set(qn('w:type'), 'dxa')
        _tbl_pr.append(_w)
        # Layout fixed
        _layout = OxmlElement('w:tblLayout')
        _layout.set(qn('w:type'), 'fixed')
        for _old_l in _tbl_pr.findall(qn('w:tblLayout')):
            _tbl_pr.remove(_old_l)
        _tbl_pr.append(_layout)

    # keep_with_next — başlık bir sonraki içerikle aynı sayfada kalsın
    if p:
        p.paragraph_format.keep_with_next = True

    return p


def _bos_satir(doc, count=1):
    """Boş satır(lar) ekle."""
    for _ in range(count):
        doc.add_paragraph()


def _sayfa_sonu(doc):
    """Sayfa sonu ekle."""
    doc.add_page_break()


def _hucre_shading(cell, hex_color):
    """Hücre arka plan rengini ayarla. hex_color: 'RRGGBB' veya '#RRGGBB'."""
    hex_color = hex_color.lstrip('#')
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _hucre_kenarsiz(cell):
    """Hücre kenarlarını kaldır."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Mevcut tcBorders varsa kaldır (XML stacking önleme)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _tablo_kenarsiz(table):
    """Tüm tablonun kenarlarını kaldır."""
    for row in table.rows:
        for cell in row.cells:
            _hucre_kenarsiz(cell)


# ═══════════════════════════════════════════════════════════
# DUAL-COLUMN LAYOUT (Kenarlıksız Tablo Grid)
# ═══════════════════════════════════════════════════════════

def iki_sutun_grid(doc, sol_genislik_cm=10.5, sag_genislik_cm=6.5, gap_cm=0.4):
    """Kenarlıksız tablo ile iki sütunlu layout oluştur.

    v4.11: gap_cm parametresi — sol hücrenin sağ kenarına internal margin olarak uygulanır.
    Sütunlar arası görsel boşluk sağlar (kapak sayfası ile tutarlı).

    Returns: (table, sol_cell, sag_cell) — içerik hücrelere yazılır.
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False

    from docx.oxml import OxmlElement

    # ── Sütun genişlikleri (python-docx API + XML seviyesi) ──
    table.columns[0].width = Cm(sol_genislik_cm)
    table.columns[1].width = Cm(sag_genislik_cm)

    tbl_pr = table._tbl.tblPr

    # v4.8: Mevcut tblW ve tblLayout'u SİL, sonra yenilerini ekle
    # python-docx varsayılan tblW type="auto" oluşturur — override etmeliyiz
    for tag in ('w:tblW', 'w:tblLayout'):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)

    # tblLayout FIXED — Word sütun genişliklerini değiştiremesin
    tbl_layout = OxmlElement('w:tblLayout')
    tbl_layout.set(qn('w:type'), 'fixed')
    tbl_pr.append(tbl_layout)

    # Toplam tablo genişliği (twips: 1 cm ≈ 567 twips)
    toplam_cm = sol_genislik_cm + sag_genislik_cm
    tbl_w = OxmlElement('w:tblW')
    tbl_w.set(qn('w:w'), str(int(toplam_cm * 567)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_w)

    sol = table.cell(0, 0)
    sag = table.cell(0, 1)

    # ── Hücre genişliklerini XML seviyesinde de zorla (tcW) ──
    for cell, cm_val in [(sol, sol_genislik_cm), (sag, sag_genislik_cm)]:
        tc_pr = cell._tc.get_or_add_tcPr()
        existing_tcw = tc_pr.find(qn('w:tcW'))
        if existing_tcw is not None:
            tc_pr.remove(existing_tcw)
        tc_w = OxmlElement('w:tcW')
        tc_w.set(qn('w:w'), str(int(cm_val * 567)))
        tc_w.set(qn('w:type'), 'dxa')
        tc_pr.append(tc_w)

    # Kenarları kaldır
    _tablo_kenarsiz(table)

    # Hücre margin'larını sıfırla (outer layout grid — margin YOK)
    _tablo_cell_margin_sifirla(table, top=0, bottom=0, left=0, right=0)

    # v4.11: Sol hücrenin sağ kenarına gap uygula — sütunlar arası görsel boşluk
    if gap_cm > 0:
        _gap_twips = str(int(gap_cm * 567))
        _tcp_sol = sol._tc.get_or_add_tcPr()
        # Mevcut tcMar varsa kaldır (idempotent)
        for old_tcm in _tcp_sol.findall(qn('w:tcMar')):
            _tcp_sol.remove(old_tcm)
        _tcm = OxmlElement('w:tcMar')
        _mar_r = OxmlElement('w:right')
        _mar_r.set(qn('w:w'), _gap_twips)
        _mar_r.set(qn('w:type'), 'dxa')
        _tcm.append(_mar_r)
        _tcp_sol.append(_tcm)

    # Satırın sayfa kırılımında bölünmesini engelle
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))

    # Dikey hizalama: üst
    sol.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    sag.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    return table, sol, sag


def cift_grafik(doc, grafik_sol_yol, grafik_sag_yol,
                baslik_sol='', baslik_sag='', kaynak='', numarala=True):
    """İki grafiği yan yana yerleştir (2'li grid).

    Grafik genişliği hücre genişliğine otomatik sığdırılır (margin payı bırakılır).
    """
    sol_cm = 8.4
    sag_cm = 8.4
    table, sol, sag = iki_sutun_grid(doc, sol_genislik_cm=sol_cm, sag_genislik_cm=sag_cm)

    _gap_bar = 0.15  # v4.13: Başlık barları arası görsel gap (cm, her taraftan)
    for idx, (cell, yol, baslik, cell_cm) in enumerate(
            [(sol, grafik_sol_yol, baslik_sol, sol_cm),
             (sag, grafik_sag_yol, baslik_sag, sag_cm)]):
        if baslik:
            if numarala:
                no = _sayac.sonraki_grafik()
                baslik = f'Grafik {no} — {baslik}'
            # v4.11: Standart başlık stili — gri arka plan + koyu metin
            p_b = _baslik_bar(cell, baslik)
            # v4.13: cift_grafik başlık barları arası gap
            # Başlık bar tablosunun genişliğini kısıtla → iki bar arasında beyaz boşluk oluşur
            _tables_in_cell = cell.tables
            if _tables_in_cell:
                _bar_tbl = _tables_in_cell[-1]
                _bar_width_cm = cell_cm - _gap_bar
                _bar_tbl.autofit = False
                _bar_tbl.allow_autofit = False
                _bar_tbl_pr = _bar_tbl._tbl.tblPr
                # Genişlik ayarla
                _existing_w = _bar_tbl_pr.find(qn('w:tblW'))
                if _existing_w is not None:
                    _bar_tbl_pr.remove(_existing_w)
                _tw = OxmlElement('w:tblW')
                _tw.set(qn('w:w'), str(int(_bar_width_cm * 567)))
                _tw.set(qn('w:type'), 'dxa')
                _bar_tbl_pr.append(_tw)
                # Bar hücre genişliği de ayarla
                _bc = _bar_tbl.cell(0, 0)
                _bc_pr = _bc._tc.get_or_add_tcPr()
                _existing_tcw = _bc_pr.find(qn('w:tcW'))
                if _existing_tcw is not None:
                    _bc_pr.remove(_existing_tcw)
                _tcw = OxmlElement('w:tcW')
                _tcw.set(qn('w:w'), str(int(_bar_width_cm * 567)))
                _tcw.set(qn('w:type'), 'dxa')
                _bc_pr.append(_tcw)
        if os.path.exists(yol):
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            # Grafik genişliği = hücre genişliği - margin (0.3cm)
            r.add_picture(yol, width=Cm(cell_cm - 0.3))
        else:
            _paragraf(cell, f'[Grafik bulunamadı: {yol}]', italic=True, color=BBB.GRI)

    # v4.11: Grafik kaynak notu YAZILMAZ — PNG içinde zaten var

    return table


# ═══════════════════════════════════════════════════════════
# HEADER / FOOTER
# ═══════════════════════════════════════════════════════════

def header_footer_ekle(doc, sol_metin, sag_metin='', logo_yolu=None):
    """Her sayfada header ve footer ekle.

    Header sol: Logo (opsiyonel) + "[Şirket] — [Rapor Türü]"
    Header separator: 2pt tema rengi çizgi
    Footer sol: Rapor adı | Footer sağ: "Sayfa X / Y"
    """
    # ── Tab stop helper — mevcut tabs'ı temizle, sağa hizalı tab ekle ──
    def _sag_tab_ekle(paragraph, pos_cm):
        """Paragrafta sağ tab stop oluştur.

        Word'ün "Header"/"Footer" paragraph stillerinin varsayılan tab stop'ları
        (center@4680tw, right@9360tw) miras yoluyla uygulanır ve bizim custom
        tab stop'umuzu gölgeler. Bu nedenle önce bu varsayılan pozisyonları
        w:tab val="clear" ile açıkça temizliyoruz.
        """
        pPr = paragraph._element.get_or_add_pPr()
        # Mevcut paragraf-düzeyi tabs'ı temizle
        for old_tabs in pPr.findall(qn('w:tabs')):
            pPr.remove(old_tabs)
        # Ayrıca pStyle'ı kaldır — stil düzeyindeki tab stop mirasını kes
        for old_pStyle in pPr.findall(qn('w:pStyle')):
            pPr.remove(old_pStyle)
        tabs = OxmlElement('w:tabs')
        # Varsayılan Word stil tab stop'larını temizle (Letter/A4 ortak pozisyonlar)
        for clear_pos in ['4680', '9360']:
            clear_tab = OxmlElement('w:tab')
            clear_tab.set(qn('w:val'), 'clear')
            clear_tab.set(qn('w:pos'), clear_pos)
            tabs.append(clear_tab)
        # Bizim sağ tab stop'umuz
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), str(int(pos_cm * 567)))  # cm → twips
        tabs.append(tab)
        pPr.append(tabs)

    # ── Tab karakteri XML olarak ekle (python \t güvenilmez) ──
    def _tab_run_ekle(paragraph):
        """Paragrafa açık w:tab XML elementi ekle."""
        r = OxmlElement('w:r')
        r.append(OxmlElement('w:tab'))
        paragraph._element.append(r)

    for section in doc.sections:
        _sayfa_cm = BBB.SAYFA_GENISLIK_CM

        # ── HEADER ──
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.clear()

        # Sağ tab stop
        _sag_tab_ekle(hp, _sayfa_cm)

        # Logo (opsiyonel)
        if logo_yolu and os.path.exists(logo_yolu):
            logo_run = hp.add_run()
            logo_run.add_picture(logo_yolu, height=Cm(0.8))
            hp.add_run('  ')

        # Sol metin
        run = hp.add_run(sol_metin)
        _font(run, size=BBB.KUCUK_PT, color=BBB.IKINCIL)

        # Tab + sağ metin
        if sag_metin:
            _tab_run_ekle(hp)
            run2 = hp.add_run(sag_metin)
            _font(run2, size=BBB.KUCUK_PT, color=BBB.GRI)

        # Separator çizgi
        _hp_pPr = hp._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), BBB.BIRINCIL_HEX)
        pBdr.append(bottom)
        _hp_pPr.append(pBdr)

        # ── FOOTER ──
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fp.clear()

        # Sağ tab stop
        _sag_tab_ekle(fp, _sayfa_cm)

        # Üst çizgi
        _fp_pPr = fp._element.get_or_add_pPr()
        fp_bdr = OxmlElement('w:pBdr')
        fp_top = OxmlElement('w:top')
        fp_top.set(qn('w:val'), 'single')
        fp_top.set(qn('w:sz'), '4')
        fp_top.set(qn('w:space'), '4')
        fp_top.set(qn('w:color'), 'CCCCCC')
        fp_bdr.append(fp_top)
        _fp_pPr.append(fp_bdr)

        # Sol metin
        if sol_metin:
            run_left = fp.add_run(sol_metin)
            _font(run_left, size=BBB.KUCUK_PT, color=BBB.GRI)

        # Tab + sağ "Sayfa X / Y"
        _tab_run_ekle(fp)
        run_pre = fp.add_run('Sayfa ')
        _font(run_pre, size=BBB.KUCUK_PT, color=BBB.GRI)

        fld_page = parse_xml(
            '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        )
        fp._element.append(fld_page)

        run_sep = fp.add_run(' / ')
        _font(run_sep, size=BBB.KUCUK_PT, color=BBB.GRI)

        fld_total = parse_xml(
            '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' w:instr=" NUMPAGES "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        )
        fp._element.append(fld_total)


# ═══════════════════════════════════════════════════════════
# İÇİNDEKİLER (TOC Field)
# ═══════════════════════════════════════════════════════════

def icindekiler_ekle(doc, bolumler=None):
    """Statik içindekiler tablosu oluştur.

    v4.13: Sayfa numaraları + dot leader desteği.
    bolumler: [(baslik, level), ...] veya [(baslik, level, sayfa_no), ...] listesi.
    Eğer None ise, assembly'den geçilmesi beklenir.
    TOC field kullanılmaz — PDF'e yazdırırken sorun çıkarmaz.
    """
    # v4.11: İçindekiler her zaman yeni sayfada — kapak sonrası explicit page break
    _baslik(doc, 'İçindekiler', level=1, sayfa_sonu=True)

    if not bolumler:
        _paragraf(doc, '[İçindekiler assembly tarafından sağlanacak]',
                  italic=True, color=BBB.GRI)
        return

    # v4.13: Tab stop pozisyonu (sayfa genişliği - sağ margin = içerik alanı)
    _toc_tab_pos_twips = int(BBB.SAYFA_GENISLIK_CM * 567)  # Sağ kenar tab stop

    for entry in bolumler:
        if len(entry) == 3:
            baslik, level, sayfa_no = entry
        else:
            baslik, level = entry
            sayfa_no = None

        p = doc.add_paragraph()
        indent_cm = 0.0 if level == 1 else 0.5 if level == 2 else 1.0
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)

        # Level 1: bold + turuncu, Level 2: normal + koyu gri
        if level == 1:
            run = p.add_run(baslik)
            _font(run, size=BBB.METIN_PT, bold=True, color=BBB.BIRINCIL)
        else:
            run = p.add_run(baslik)
            _font(run, size=BBB.METIN_PT - 1, color=BBB.IKINCIL)

        # v4.13: Dot leader + sayfa numarası
        if sayfa_no is not None:
            # Tab stop ayarla — sağ hizalı, dot leader
            pPr = p._element.get_or_add_pPr()
            # Mevcut tabs temizle
            for old_tabs in pPr.findall(qn('w:tabs')):
                pPr.remove(old_tabs)
            tabs = OxmlElement('w:tabs')
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:leader'), 'dot')  # Nokta dolgu (dot leader)
            # indent'e göre tab pozisyonu ayarla
            _effective_tab = _toc_tab_pos_twips - int(indent_cm * 567)
            tab.set(qn('w:pos'), str(_effective_tab))
            tabs.append(tab)
            pPr.append(tabs)

            # Tab karakteri + sayfa numarası ekle
            tab_run = OxmlElement('w:r')
            tab_run.append(OxmlElement('w:tab'))
            p._element.append(tab_run)

            # Sayfa numarası
            pg_run = p.add_run(str(sayfa_no))
            if level == 1:
                _font(pg_run, size=BBB.METIN_PT, bold=True, color=BBB.BIRINCIL)
            else:
                _font(pg_run, size=BBB.METIN_PT - 1, color=BBB.IKINCIL)


# ═══════════════════════════════════════════════════════════
# TABLO ÜRETİMİ (Profesyonel)
# ═══════════════════════════════════════════════════════════

def _hucre_border_modern(cell, top=True, bottom=True, left=False, right=False,
                         color='E0E0E0', sz='4'):
    """Modern tablo kenarları: sadece üst/alt ince çizgi. color: 'RRGGBB' veya '#RRGGBB'. sz: int veya str."""
    color = str(color).lstrip('#')
    sz = str(sz)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Mevcut tcBorders varsa kaldır (XML stacking önleme)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for border_name, enabled in [('top', top), ('bottom', bottom),
                                  ('left', left), ('right', right)]:
        border = OxmlElement(f'w:{border_name}')
        if enabled:
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
        else:
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _tablo_border_sifirla(table):
    """Tablo seviyesinde TÜM border'ları kaldır (Table Grid stili dahil).

    Word varsayılan olarak 'Table Grid' stili uygular ve tüm hücrelere border ekler.
    Bu fonksiyon table-level tblBorders'ı 'none' yaparak bu davranışı engeller.
    """
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    # Mevcut tblBorders varsa kaldır
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def tablo_stil(table, baslik_renk=True):
    """BBB stili: turuncu başlık + koyu metin + temiz beyaz satırlar + çizgisiz.

    AK Yatırım referans: hiçbir çizgi yok, sadece başlık arka plan rengi ile ayrım.
    """
    # Önce tüm table-level ve cell-level border'ları kaldır
    _tablo_border_sifirla(table)

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            # Tüm hücrelerin border'larını açıkça kaldır
            _hucre_border_modern(cell, top=False, bottom=False, left=False, right=False)
            if i == 0 and baslik_renk:
                # Başlık: turuncu arka plan + koyu bold metin
                _hucre_shading(cell, BBB.BIRINCIL_HEX)
                for p in cell.paragraphs:
                    for run in p.runs:
                        _font(run, size=BBB.TABLO_PT, bold=True, color=BBB.SIYAH)


def _hucre_deger_renk(cell, value_str):
    """Hücre değerine göre koşullu renklendirme (yeşil/kırmızı).

    Kurallar:
      - '+' ile başlayan veya pozitif % → yeşil (BBB.OLUMLU)
      - '-' ile başlayan veya negatif % → kırmızı (BBB.OLUMSUZ)
      - 'BEAT', 'artış', 'Güçlü' → yeşil
      - 'MISS', 'düşüş', 'Zayıf' → kırmızı
    """
    s = value_str.strip()
    if not s or s == '—':
        return None

    # Pozitif/negatif sayılar
    if s.startswith('+') or s.startswith('▲'):
        return BBB.OLUMLU
    if s.startswith('-') and len(s) > 1 and (s[1].isdigit() or s[1] == '%'):
        return BBB.OLUMSUZ
    if s.startswith('▼'):
        return BBB.OLUMSUZ

    # Yüzde değerler
    pct_match = re.match(r'^[%]?\s*([+-]?\d)', s) or re.search(r'([+-]\d[^)]*%)$', s)
    if pct_match:
        try:
            sign_char = pct_match.group(1)[0] if pct_match.group(1) else None
            if sign_char == '+':
                return BBB.OLUMLU
            elif sign_char == '-':
                return BBB.OLUMSUZ
        except (IndexError, ValueError):
            pass

    # Anahtar kelimeler
    s_lower = s.lower()
    for kw in ('beat', 'artış', 'güçlü', 'olumlu', 'yüksek potansiyel'):
        if kw in s_lower:
            return BBB.OLUMLU
    for kw in ('miss', 'düşüş', 'zayıf', 'olumsuz'):
        if kw in s_lower:
            return BBB.OLUMSUZ

    return None


def _tablo_cell_margin_sifirla(table, top=0, bottom=0, left=30, right=30):
    """Tablo seviyesinde varsayılan hücre margin'larını sıfırla.

    Tüm hücrelere etki eder — cell-level override gerekmez.
    Değerler twips (1/20 pt) cinsinden. 30 twips ≈ 0.5mm.
    """
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    # Mevcut tblCellMar varsa kaldır
    existing = tblPr.find(qn('w:tblCellMar'))
    if existing is not None:
        tblPr.remove(existing)
    tcMar = OxmlElement('w:tblCellMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tblPr.append(tcMar)


def _turkce_duzelt(text):
    """Yaygın Türkçe karakter eksikliklerini düzelt.

    Assembly'den gelen metinlerde sıkça eksik kalan ş, ç, ğ, ı, ö, ü, İ
    harflerini kelime bazlı düzeltir. Tablo başlıkları ve veri hücrelerinde
    otomatik uygulanır.
    """
    # Kelime bazlı düzeltme sözlüğü (küçük harf → doğru yazım)
    _corrections = {
        'satis': 'Satış', 'Satis': 'Satış', 'satıs': 'Satış',
        'hasılat': 'Hasılat', 'hasilat': 'Hasılat',
        'buyume': 'Büyüme', 'Buyume': 'Büyüme', 'büyume': 'Büyüme',
        'degerlendirme': 'Değerlendirme', 'Degerlendirme': 'Değerlendirme',
        'aciklama': 'Açıklama', 'Aciklama': 'Açıklama',
        'deger': 'Değer', 'Deger': 'Değer',
        'carpan': 'Çarpan', 'Carpan': 'Çarpan',
        'yontem': 'Yöntem', 'Yontem': 'Yöntem',
        'senaryo': 'Senaryo',
        'sirket': 'Şirket', 'Sirket': 'Şirket',
        'ulke': 'Ülke', 'Ulke': 'Ülke',
        'olasilik': 'Olasılık', 'Olasilik': 'Olasılık',
        'gerekce': 'Gerekçe', 'Gerekce': 'Gerekçe',
        'dusuk': 'Düşük', 'Dusuk': 'Düşük',
        'yuksek': 'Yüksek', 'Yuksek': 'Yüksek',
        'iyimser': 'İyimser', 'Iyimser': 'İyimser',
        'kotumser': 'Kötümser', 'Kotumser': 'Kötümser',
        'agirlik': 'Ağırlık', 'Agirlik': 'Ağırlık',
        'sutun': 'Sütun', 'Sutun': 'Sütun',
        'olcum': 'Ölçüm', 'Olcum': 'Ölçüm',
        'yukselis': 'Yükseliş', 'Yukselis': 'Yükseliş',
        'dususu': 'Düşüşü', 'Dususu': 'Düşüşü',
        'dusus': 'Düşüş', 'Dusus': 'Düşüş',
        'iyilesme': 'İyileşme', 'Iyilesme': 'İyileşme',
        'iyilesen': 'İyileşen', 'Iyilesen': 'İyileşen',
        'genisleme': 'Genişleme', 'genislemesi': 'Genişlemesi',
        'calisan': 'Çalışan', 'Calisan': 'Çalışan',
        'sektor': 'Sektör', 'Sektor': 'Sektör',
        'gerceklesme': 'Gerçekleşme', 'Gerceklesme': 'Gerçekleşme',
        'sonuc': 'Sonuç', 'Sonuc': 'Sonuç',
        'marji': 'Marjı', 'Marji': 'Marjı',
        'basi': 'Başı', 'Basi': 'Başı',
        'degisim': 'Değişim', 'Degisim': 'Değişim',
        'aralik': 'Aralık', 'Aralik': 'Aralık',
        'ceyreklik': 'Çeyreklik', 'ceyrek': 'çeyrek',
        'artisi': 'Artışı', 'artis': 'Artış',
    }
    result = text
    for wrong, correct in _corrections.items():
        # Kelime sınırı ile değiştir (yanlış pozitif önleme)
        result = re.sub(r'\b' + re.escape(wrong) + r'\b', correct, result)
    return result


def _hesapla_tablo_genislik(basliklar, satirlar, doc_or_cell, kompakt=False, genislik_cm=None):
    """Tablo genişliğini content-optimal olarak hesapla (tablo oluşturmadan ÖNCE).

    v4.14: _baslik_bar ile tablo genişliğinin simetrik olması için
    genişlik hesaplaması tablo oluşturulmasından ayrıştırıldı.

    Returns: (toplam_genislik_cm, final_widths_list)
    """
    n_cols = len(basliklar)
    font_size = BBB.TABLO_PT - 1 if kompakt else BBB.TABLO_PT

    # Sayfa limiti
    if hasattr(doc_or_cell, 'width') and hasattr(doc_or_cell, 'vertical_alignment'):
        sayfa_cm = doc_or_cell.width.cm
    else:
        sayfa_cm = 15.0

    # Her sütunun max karakter uzunluğu
    max_lens = []
    for j in range(n_cols):
        col_max = len(str(basliklar[j]))
        for satir in satirlar:
            if j < len(satir):
                col_max = max(col_max, len(str(satir[j])))
        max_lens.append(col_max)

    cm_per_char = 0.22 if font_size <= 8 else 0.25
    padding_cm = 0.3
    optimal_widths = [ml * cm_per_char + padding_cm for ml in max_lens]
    toplam_optimal = sum(optimal_widths)

    if genislik_cm is not None:
        toplam_cm = min(genislik_cm, sayfa_cm)
        ratio = toplam_cm / toplam_optimal if toplam_optimal > 0 else 1.0
        final_widths = [w * ratio for w in optimal_widths]
    elif toplam_optimal <= sayfa_cm:
        final_widths = optimal_widths
    else:
        ratio = sayfa_cm / toplam_optimal
        final_widths = [w * ratio for w in optimal_widths]

    return sum(final_widths), final_widths


def tablo_ekle(doc_or_cell, basliklar, satirlar, kaynak='', kompakt=False,
               kosullu_renk=True, baslik_metin='', numarala=True,
               genislik_cm=None):
    """Profesyonel tablo oluştur.

    doc_or_cell: Document veya table cell
    kompakt: True ise daha küçük font
    kosullu_renk: True ise pozitif/negatif değerler yeşil/kırmızı
    baslik_metin: Tablo üst başlığı (varsa otomatik numaralanır)
    numarala: True ise başlık önüne "Tablo N — " eklenir
    genislik_cm: Tablo toplam genişliği (cm). None ise content-optimal hesaplanır.
                 Aynı sayfadaki tabloların genişlik uyumu için kullanılır.
    """
    target = doc_or_cell

    # v4.14: Genişlik hesaplamasını ÖNCE yap — _baslik_bar ile simetri sağlamak için
    # Türkçe düzeltme öncesi ham veri ile hesapla (karakter sayıları yaklaşık aynı)
    hesaplanan_toplam, hesaplanan_widths = _hesapla_tablo_genislik(
        basliklar, satirlar, doc_or_cell, kompakt=kompakt, genislik_cm=genislik_cm)

    if baslik_metin:
        if numarala:
            no = _sayac.sonraki_tablo()
            baslik_metin = f'Tablo {no} — {baslik_metin}'
        # v4.14: baslik_bar'a hesaplanan genişliği aktar — content-optimal durumda da simetrik
        p_tb = _baslik_bar(target, baslik_metin, genislik_cm=hesaplanan_toplam)
        # Tablo başlığı tablodan kopmasın
        if p_tb:
            p_tb.paragraph_format.keep_with_next = True
    # Türkçe karakter otomatik düzeltme (engine seviyesi)
    basliklar = [_turkce_duzelt(str(b)) for b in basliklar]
    satirlar = [[_turkce_duzelt(str(v)) for v in satir] for satir in satirlar]
    if kaynak:
        kaynak = _turkce_duzelt(kaynak)

    target = doc_or_cell
    n_cols = len(basliklar)
    n_rows = 1 + len(satirlar)
    table = target.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Table Grid'  # Önce stil ata, sonra border'ları kaldır
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False

    # v4.9: tblLayout FIXED + tblW dxa — Word'ün auto-resize yapmasını engelle
    # iki_sutun_grid ile tutarlı — iç tablolarda da genişlik XML seviyesinde zorlanır
    _tbl_pr = table._tbl.tblPr
    for _tag in ('w:tblLayout',):
        _existing = _tbl_pr.find(qn(_tag))
        if _existing is not None:
            _tbl_pr.remove(_existing)
    _tbl_layout = OxmlElement('w:tblLayout')
    _tbl_layout.set(qn('w:type'), 'fixed')
    _tbl_pr.append(_tbl_layout)

    # Zebra/banding satırlarını devre dışı bırak
    tbl_pr = table._tbl.tblPr
    tbl_look = tbl_pr.find(qn('w:tblLook'))
    if tbl_look is None:
        tbl_look = OxmlElement('w:tblLook')
        tbl_pr.append(tbl_look)
    tbl_look.set(qn('w:firstRow'), '1')
    tbl_look.set(qn('w:lastRow'), '0')
    tbl_look.set(qn('w:firstColumn'), '0')
    tbl_look.set(qn('w:lastColumn'), '0')
    tbl_look.set(qn('w:noHBand'), '1')  # Yatay banding KAPALI
    tbl_look.set(qn('w:noVBand'), '1')  # Dikey banding KAPALI

    # Tablo seviyesinde hücre margin'larını sıfırla (en etkili yöntem)
    _tablo_cell_margin_sifirla(table)

    font_size = BBB.KUCUK_PT if kompakt else BBB.TABLO_PT
    cell_line_spacing = Pt(8) if kompakt else Pt(9)

    # Başlıklar — koyu metin (tablo_stil ile uyumlu)
    for j, b in enumerate(basliklar):
        cell = table.cell(0, j)
        # İlk sütun sol hizalı, diğerleri sağ hizalı (başlık da veri ile uyumlu)
        h_align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        _paragraf(cell, str(b), size=font_size, bold=True, color=BBB.SIYAH,
                  alignment=h_align,
                  space_after=0, space_before=0, line_spacing=cell_line_spacing)

    # Veriler — ilk sütun daima sol, sayısal sütunlar daima sağ
    for i, satir in enumerate(satirlar):
        for j, val in enumerate(satir):
            cell = table.cell(i + 1, j)
            val_str = str(val)
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            deger_renk = _hucre_deger_renk(cell, val_str) if kosullu_renk and j > 0 else None
            _paragraf(cell, val_str, size=font_size, color=deger_renk or BBB.METIN,
                      bold=bool(deger_renk),
                      alignment=align, space_after=0, space_before=0,
                      line_spacing=cell_line_spacing)

    # Satır yüksekliğini ayarla — kompakt tablolar için daha sıkı
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    if kompakt:
        header_h, data_h = Pt(11), Pt(9)
    else:
        header_h, data_h = Pt(13), Pt(11)
    for idx, row in enumerate(table.rows):
        row.height = header_h if idx == 0 else data_h
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    # Tüm hücrelerin vertical alignment'ını CENTER yap
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # v4.14: Genişlik hesaplaması artık fonksiyon başında yapılıyor (_hesapla_tablo_genislik)
    # _baslik_bar ile simetri sağlamak için önceden hesaplanmış değerleri kullan
    final_widths = hesaplanan_widths

    for col_idx in range(n_cols):
        col_cm = final_widths[col_idx]
        for cell in table.columns[col_idx].cells:
            cell.width = Cm(col_cm)

    # v4.9: tblW dxa ile toplam tablo genişliğini XML'de zorla
    toplam_final = sum(final_widths)
    for _old_tw in _tbl_pr.findall(qn('w:tblW')):
        _tbl_pr.remove(_old_tw)
    _tbl_w = OxmlElement('w:tblW')
    _tbl_w.set(qn('w:w'), str(int(toplam_final * 567)))
    _tbl_w.set(qn('w:type'), 'dxa')
    _tbl_pr.append(_tbl_w)

    tablo_stil(table)

    # v4.7: Tablo başlık satırını ve ilk veri satırını birlikte tut
    # Son satıra keep_with_next KOYMA — boş sayfa taşmasına yol açar
    for idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                if idx < min(2, n_rows - 1):
                    p.paragraph_format.keep_with_next = True
                p.paragraph_format.widow_control = True

    if kaynak:
        _kaynak(doc_or_cell if hasattr(doc_or_cell, 'add_paragraph') and not hasattr(doc_or_cell, 'vertical_alignment') else doc_or_cell, kaynak)
    else:
        # v4.14: Kaynak satırı olmayan tablolarda da sonrası boşluk bırak
        # Sadece document-level tablolarda (cell-içi tablolarda değil)
        if hasattr(doc_or_cell, 'add_paragraph') and not hasattr(doc_or_cell, 'vertical_alignment'):
            _spacer = doc_or_cell.add_paragraph()
            _spacer.paragraph_format.space_before = Pt(0)
            _spacer.paragraph_format.space_after = Pt(8)
            _spacer.paragraph_format.line_spacing = Pt(2)

    return table


def _is_numeric(s):
    """Sayısal string kontrolü."""
    s = s.strip().replace(',', '').replace('.', '').replace('%', '').replace('-', '').replace('x', '')
    return s.isdigit() if s else False


# ═══════════════════════════════════════════════════════════
# GRAFİK EMBED
# ═══════════════════════════════════════════════════════════

def grafik_ekle(doc_or_cell, dosya_yolu, boyut='ORTA', baslik='', kaynak='',
                numarala=True):
    """Grafik embed. boyut: 'TAM' (15.5cm) | 'YARIM' (7.4cm) | 'GRID' (7.2cm) | 'KUCUK' (5cm)

    numarala=True ise başlık önüne otomatik "Grafik N — " eklenir.
    """
    # None path guard — assembly'de grafik bulunamazsa None dönebilir
    if dosya_yolu is None:
        _paragraf(doc_or_cell, '[Grafik yolu belirtilmedi]', italic=True, color=BBB.GRI)
        return
    target = doc_or_cell
    boyut_map = {
        'TAM': BBB.GRAFIK_TAM,
        'ORTA': BBB.GRAFIK_ORTA,
        'YARIM': BBB.GRAFIK_YARIM,
        'GRID': BBB.GRAFIK_GRID,
        'KUCUK': BBB.GRAFIK_KUCUK,
        'PASTA': BBB.GRAFIK_PASTA,
    }
    width = boyut_map.get(boyut, BBB.GRAFIK_TAM)

    if baslik:
        if numarala:
            no = _sayac.sonraki_grafik()
            baslik = f'Grafik {no} — {baslik}'
        # v4.11: Standart başlık stili — gri arka plan + koyu metin (grafik_kutucuk ile tutarlı)
        p_baslik = _baslik_bar(target, baslik)
        # Başlık grafikten kopmasın
        if p_baslik:
            p_baslik.paragraph_format.keep_with_next = True

    if not os.path.exists(dosya_yolu):
        _paragraf(target, f'[Grafik bulunamadı: {dosya_yolu}]', italic=True, color=BBB.GRI)
        return

    if hasattr(target, 'add_picture'):
        # Document
        target.add_picture(dosya_yolu, width=width)
        target.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Table cell
        p = target.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(dosya_yolu, width=width)

    # v4.11: Grafik kaynak notu YAZILMAZ — PNG'nin içinde zaten "Kaynak: ..." metni var.
    # kaynak parametresi geriye uyumluluk için kabul edilir ama ignore edilir.
    # Çift kaynak notu sorunu düzeltildi.


def grafik_kutucuk(doc_or_cell, dosya_yolu, baslik, kaynak='', boyut='ORTA',
                   numarala=True):
    """Çerçeveli grafik: başlık barı + grafik + kaynak notu.

    Profesyonel rapor standardı: grafik+başlık+kaynak tek bir birim olarak
    1×1 kenarlıklı tablo içinde sunulur. Başlık barı tema rengi arka planlı.
    """
    if dosya_yolu is None or (dosya_yolu and not os.path.exists(dosya_yolu)):
        _paragraf(doc_or_cell, f'[Grafik bulunamadı: {dosya_yolu}]',
                  italic=True, color=BBB.GRI)
        return

    target = doc_or_cell
    boyut_map = {
        'TAM': BBB.GRAFIK_TAM,
        'ORTA': BBB.GRAFIK_ORTA,
        'YARIM': BBB.GRAFIK_YARIM,
        'GRID': BBB.GRAFIK_GRID,
        'KUCUK': BBB.GRAFIK_KUCUK,
        'PASTA': BBB.GRAFIK_PASTA,
    }
    width = boyut_map.get(boyut, BBB.GRAFIK_TAM)

    # v4.11: Kaynak notu YAZILMAZ — PNG içinde zaten var. Her zaman 2 satırlı: [başlık] [grafik]
    n_rows = 2
    frame = target.add_table(rows=n_rows, cols=1)
    frame.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Çerçeve genişliği = grafik genişliği + padding
    for cell in frame.columns[0].cells:
        cell.width = width

    # ── Satır 0: Başlık barı ──
    header_cell = frame.cell(0, 0)
    _hucre_shading(header_cell, BBB.ACIK_GRI_HEX)  # Sıcak bej-gri — sayfa uyumlu nötr kutu
    if numarala:
        no = _sayac.sonraki_grafik()
        baslik = f'Grafik {no} — {baslik}'
    _paragraf(header_cell, baslik, size=BBB.TABLO_PT, bold=True, color=BBB.IKINCIL,
              space_after=2, space_before=2, alignment=WD_ALIGN_PARAGRAPH.LEFT,
              line_spacing=Pt(10))
    header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    frame.rows[0].height = Pt(18)

    # ── Satır 1: Grafik ──
    chart_cell = frame.cell(1, 0)
    p = chart_cell.paragraphs[0]  # Varsayılan paragrafı yeniden kullan
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(dosya_yolu, width=width)

    # v4.11: Kaynak notu satırı kaldırıldı — PNG içinde zaten var

    # Çerçeve: ince border + orphan koruması
    for row_idx, row in enumerate(frame.rows):
        for cell in row.cells:
            _hucre_border_modern(cell, top=True, bottom=True, left=True, right=True,
                                 color='CCCCCC', sz='4')
            # v4.7: Başlık + grafik satırlarını birlikte tut, son satır (kaynak) hariç
            if row_idx < n_rows - 1:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True

    return frame


def grafik_grid(doc, grafikler, cols=2, rows=None, baslik=''):
    """NxM grafik grid — tek sayfada birden çok grafik.

    grafikler: [{'yol': '...', 'baslik': '...', 'kaynak': '...'}, ...]
    cols: sütun sayısı (varsayılan 2)
    rows: satır sayısı (None ise otomatik hesaplanır)
    baslik: Grid üst başlığı (opsiyonel)
    """
    import math
    n = len(grafikler)
    if n == 0:
        return
    if rows is None:
        rows = math.ceil(n / cols)

    if baslik:
        _baslik(doc, baslik, level=2)

    # Kenarlıksız layout tablo
    grid = doc.add_table(rows=rows, cols=cols)
    grid.autofit = False
    grid.allow_autofit = False
    _tablo_kenarsiz(grid)

    # Her hücrenin genişliği
    cell_cm = BBB.SAYFA_GENISLIK_CM / cols
    for col in grid.columns:
        for cell in col.cells:
            cell.width = Cm(cell_cm)

    # Grafik boyutu: hücre genişliği - padding
    grafik_w = Cm(cell_cm - 0.6)

    for idx, g in enumerate(grafikler):
        r = idx // cols
        c = idx % cols
        if r >= rows:
            break

        cell = grid.cell(r, c)
        # cell.text = '' KULLANMA — ghost paragraph yaratır (Normal stil miras alır)
        # Varsayılan paragrafı temizle
        for _p in cell.paragraphs:
            for _r in list(_p.runs):
                _r._element.getparent().remove(_r._element)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Mini başlık
        g_baslik = g.get('baslik', '')
        if g_baslik:
            _paragraf(cell, g_baslik, size=BBB.KUCUK_PT, bold=True,
                      color=BBB.BIRINCIL, space_after=2, space_before=2,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=Pt(10))

        # Grafik embed
        yol = g.get('yol', '')
        if yol and os.path.exists(yol):
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run()
            run.add_picture(yol, width=grafik_w)
        else:
            _paragraf(cell, f'[Grafik: {yol}]', italic=True, color=BBB.GRI,
                      size=BBB.KUCUK_PT)

        # Kaynak notu
        g_kaynak = g.get('kaynak', '')
        if g_kaynak:
            _paragraf(cell, g_kaynak, size=7, italic=True, color=BBB.GRI,
                      space_after=1, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      line_spacing=Pt(9))

    return grid


def metin_grafik_layout(doc, metin, grafik_yolu, grafik_baslik='',
                        grafik_kaynak='', metin_sol=True,
                        metin_genislik_cm=10.0, grafik_genislik_cm=7.0,
                        numarala=True, max_paragraf=None):
    """Dual-column içerik layout: metin + grafik yan yana.

    AK Yatırım / Deniz Yatırım tarzı: bir tarafta analiz metni,
    diğer tarafta destekleyici grafik veya tablo.

    metin: str veya list[str] — paragraf(lar)
    grafik_yolu: grafik dosya yolu
    metin_sol: True ise sol metin + sağ grafik (varsayılan)

    v4.13: max_paragraf parametresi — dual-panel'e yerleştirilecek maksimum paragraf sayısı.
    Aşan paragraflar layout dışında normal tam genişlik paragraf olarak devam eder.
    None ise tüm paragraflar panele yerleştirilir (eski davranış).
    """
    if metin_sol:
        table, sol, sag = iki_sutun_grid(doc, sol_genislik_cm=metin_genislik_cm,
                                          sag_genislik_cm=grafik_genislik_cm)
        metin_cell, grafik_cell = sol, sag
    else:
        table, sol, sag = iki_sutun_grid(doc, sol_genislik_cm=grafik_genislik_cm,
                                          sag_genislik_cm=metin_genislik_cm)
        grafik_cell, metin_cell = sol, sag

    # ── Metin bölümü ──
    if isinstance(metin, str):
        paragraflar = [metin]
    else:
        paragraflar = list(metin)

    # v4.13: max_paragraf ile taşma kontrolü
    panel_paragraflar = paragraflar if max_paragraf is None else paragraflar[:max_paragraf]
    tasan_paragraflar = [] if max_paragraf is None else paragraflar[max_paragraf:]

    for para_text in panel_paragraflar:
        _paragraf(metin_cell, para_text, size=BBB.METIN_PT, color=BBB.METIN,
                  space_after=3, line_spacing=Pt(11))

    # ── Grafik bölümü (çerçeveli kutucuk) ──
    grafik_w = Cm(grafik_genislik_cm - 0.5)  # Hücre içi padding

    if grafik_baslik:
        if numarala:
            no = _sayac.sonraki_grafik()
            grafik_baslik = f'Grafik {no} — {grafik_baslik}'
        # v4.11: Standart başlık stili — gri arka plan + koyu metin
        p_gb = _baslik_bar(grafik_cell, grafik_baslik)

    if grafik_yolu and os.path.exists(grafik_yolu):
        p = grafik_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(grafik_yolu, width=grafik_w)
    else:
        _paragraf(grafik_cell, f'[Grafik: {grafik_yolu}]',
                  italic=True, color=BBB.GRI, size=BBB.KUCUK_PT)

    # v4.11: Grafik kaynak notu YAZILMAZ — PNG içinde zaten var

    # v4.13: Taşan paragraflar — layout dışında normal tam genişlik devam
    for para_text in tasan_paragraflar:
        _paragraf(doc, para_text, size=BBB.METIN_PT, color=BBB.METIN,
                  space_after=4, line_spacing=Pt(12))

    return table


# ═══════════════════════════════════════════════════════════
# SAYFA 1: AK YATIRIM TARZI KAPAK
# ═══════════════════════════════════════════════════════════

def kapak_sayfasi(doc, ticker, sirket_adi='', rapor_turu='KAPSAM BAŞLATMA',
                  tarih='', sektor_kategori='',
                  headline='',
                  tavsiye='TUT',
                  mevcut_fiyat='—', hedef_fiyat='—', potansiyel='—',
                  pd='—', ev='—', hafta52='—',
                  halka_aciklik='—', gunluk_hacim='—',
                  analist='Kaya (AI Agent)', analist_kurum='borsadabibasina.com',
                  tez_pillars=None,
                  degerleme_ozeti='', risk_ozeti='',
                  fiyat_bilgileri=None,
                  finansal_ozet=None,
                  carpanlar=None,
                  ortaklik_yapisi=None,
                  grafik_yolu='',
                  alt_notlar=None):
    """Top-tier kapak sayfası v2.1 — AK Yatırım / Goldman Sachs referans.

    SOL PANEL (~60%): Sektör satırı + şirket adı (22pt) + headline + detaylı tez
                      pillar'ları (3-5 cümle/pillar) + değerleme özeti + risk cümlesi.
                      Hedef: Sayfa sonuna kadar dolduracak zenginlikte metin.

    SAĞ PANEL (~40%): Beyaz arka plan. Üstte analist bilgisi. Yapılandırılmış tablolar:
                       (1) Fiyat Bilgileri key-value tablo
                       (2) Finansal Özet tablo (12+ satır)
                       (3) Çarpanlar tablo
                       (4) Fiyat performans grafiği (en altta)
                       (5) Alt notlar (grafiğin altında, sağ panel içinde)

    Args:
        sektor_kategori: "Perakende / Bebek & Anne" — üst satır
        headline: "Konsolidasyon Lideri ile Büyüme Hikayesi" — italic catchline
        tez_pillars: [{'baslik': '...', 'metin': '...'}, ...] — 3-5 cümle/pillar
        degerleme_ozeti: Pillar'lardan sonra kısa değerleme cümlesi
        risk_ozeti: Değerleme altında kısa risk cümlesi
        fiyat_bilgileri: [('Etiket', 'Değer'), ...] — sağ panel üst key-value
        finansal_ozet: {'basliklar': [...], 'satirlar': [[...], ...]}
        carpanlar: {'basliklar': [...], 'satirlar': [[...], ...]}
        ortaklik_yapisi: {'basliklar': [...], 'satirlar': [[...], ...]} — ortaklık tablosu
        alt_notlar: ['*Not 1', '*Not 2'] — sağ panel alt notları
    """
    baslik = sirket_adi or ticker
    _genislik = BBB.SAYFA_GENISLIK_CM

    # ─── ÜST BANT: Sektör/Kategori + Tarih (sol-sağ hizalı) ───
    if sektor_kategori or tarih:
        p_ust = doc.add_paragraph()
        p_ust.paragraph_format.space_after = Pt(2)
        p_ust.paragraph_format.space_before = Pt(0)
        # Sektör — Tarih formatı (dash separator, tab boşluğu yok)
        _ust_parts = []
        if sektor_kategori:
            _ust_parts.append(sektor_kategori)
        if tarih:
            _ust_parts.append(tarih)
        run_s = p_ust.add_run(' — '.join(_ust_parts))
        _font(run_s, size=BBB.KUCUK_PT, color=BBB.GRI)

    # ─── ŞİRKET ADI (22pt bold) + RAPOR TÜRÜ ───
    p_ad = doc.add_paragraph()
    p_ad.paragraph_format.space_after = Pt(1)
    p_ad.paragraph_format.space_before = Pt(0)
    run_ad = p_ad.add_run(f'{baslik} ({ticker})')
    _font(run_ad, size=22, bold=True, color=BBB.BIRINCIL)

    p_tur = doc.add_paragraph()
    p_tur.paragraph_format.space_after = Pt(2)
    p_tur.paragraph_format.space_before = Pt(0)
    run_tur = p_tur.add_run(rapor_turu)
    _font(run_tur, size=10, bold=True, color=BBB.IKINCIL)

    # ─── HEADLINE (italic catchline) ───
    if headline:
        p_hl = doc.add_paragraph()
        p_hl.paragraph_format.space_after = Pt(2)   # v4.12: 4→2pt — S.1 sığma
        p_hl.paragraph_format.space_before = Pt(0)
        run_hl = p_hl.add_run(headline)
        _font(run_hl, size=11, bold=True, italic=True, color=BBB.METIN)

    # ─── DUAL-COLUMN: Sol tez metni (~60%) | Sağ tablolar (~40%) — AK Yatırım referans ───
    # v4.11: Gap toplam genişlikten değil sol panelden düşülür — sağ kenar boşluk fix
    _gap_cm = 0.4  # paneller arası görsel boşluk (sol hücrenin sağ margin'ı olarak uygulanır)
    sol_cm = round(_genislik * 0.60, 1)   # 11.0 cm (gap dahil — görünen metin alanı = 11.0 - 0.4 = 10.6)
    sag_cm = round(_genislik * 0.40, 1)   # 7.3 cm (tam genişlik — sağ kenarda boşluk kalmaz)
    table, sol, sag = iki_sutun_grid(doc, sol_genislik_cm=sol_cm, sag_genislik_cm=sag_cm,
                                      gap_cm=_gap_cm)
    # v4.12: cantSplit kaldır — kapak tablosu page-filling layout, tek sayfaya sığmaması
    # normal davranış. cantSplit kalırsa tüm tablo S.2'ye taşınır (S.1 boş kalır).
    _tr = table.rows[0]._tr
    _trPr = _tr.find(qn('w:trPr'))
    if _trPr is not None:
        for _cs in _trPr.findall(qn('w:cantSplit')):
            _trPr.remove(_cs)

    # ═══ SOL PANEL: Detaylı Tez Metni (iki yana yaslı) ═══
    # Hedef: Sayfa sonuna kadar dolduracak zenginlikte (min 400 kelime)
    if tez_pillars:
        for pillar in tez_pillars:
            p = sol.add_paragraph()
            p.paragraph_format.space_after = Pt(5)   # v4.12: 7→5pt — S.1 sığma optimizasyonu
            p.paragraph_format.space_before = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # v4.8: iki yana yaslı
            # ■ Bold başlık + metin
            run = p.add_run(f'■ {pillar.get("baslik", "")}. ')
            _font(run, size=BBB.METIN_PT, bold=True, color=BBB.METIN)
            run2 = p.add_run(pillar.get('metin', ''))
            _font(run2, size=BBB.METIN_PT, color=BBB.METIN)
    else:
        _paragraf(sol, '[Tez pillarları buraya — tez_pillars parametresi]',
                  italic=True, color=BBB.GRI)

    # Değerleme özeti (pillar'lardan sonra)
    if degerleme_ozeti:
        p_do = sol.add_paragraph()
        p_do.paragraph_format.space_after = Pt(2)   # v4.12: 4→2pt — S.1 sığma
        p_do.paragraph_format.space_before = Pt(1)
        p_do.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_dl = p_do.add_run('Değerleme: ')
        _font(run_dl, size=BBB.METIN_PT, bold=True, color=BBB.BIRINCIL)
        run_dt = p_do.add_run(degerleme_ozeti)
        _font(run_dt, size=BBB.METIN_PT, color=BBB.METIN)

    # Risk özeti
    if risk_ozeti:
        p_ro = sol.add_paragraph()
        p_ro.paragraph_format.space_after = Pt(2)
        p_ro.paragraph_format.space_before = Pt(0)
        p_ro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_rl = p_ro.add_run('Temel riskler: ')
        _font(run_rl, size=BBB.METIN_PT, bold=True, color=BBB.OLUMSUZ)
        run_rt = p_ro.add_run(risk_ozeti)
        _font(run_rt, size=BBB.METIN_PT, color=BBB.METIN)

    # ═══ SAĞ PANEL: Yapılandırılmış Tablolar (beyaz arka plan) ═══
    # v4.8: Arka plan shading kaldırıldı — temiz beyaz (referans standart)

    # ─── Analist bilgisi (sağ panelin EN ÜSTÜnde, sağa yaslı) ───
    p_analist = sag.add_paragraph()
    p_analist.paragraph_format.space_before = Pt(0)
    p_analist.paragraph_format.space_after = Pt(3)
    p_analist.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_an = p_analist.add_run(f'Analist: {analist} | ')
    _font(run_an, size=BBB.KUCUK_PT, italic=True, color=BBB.GRI)
    # Kurum adı tıklanabilir link
    _analist_url = analist_kurum if analist_kurum.startswith('http') else ''
    _analist_display = analist_kurum.replace('https://', '').replace('http://', '').rstrip('/')
    if '/' in _analist_display:
        _analist_display = _analist_display.split('/')[0]  # sadece domain
    if _analist_url:
        add_hyperlink(p_analist, _analist_display, _analist_url,
                      font_size=BBB.KUCUK_PT)
    else:
        run_k = p_analist.add_run(_analist_display)
        _font(run_k, size=BBB.KUCUK_PT, italic=True, color=BBB.GRI)

    # ─── (1) Piyasa Bilgileri — key-value tablo (diğer tablolarla aynı stil) ───
    if fiyat_bilgileri is None:
        fiyat_bilgileri = [
            ('Tavsiye', str(tavsiye)),
            ('Adil Değer Tahmini', f'{hedef_fiyat} TL'),
            ('Mevcut Fiyat', f'{mevcut_fiyat} TL'),
            ('Getiri Potansiyeli', str(potansiyel)),
            ('Piyasa Değeri', str(pd)),
            ('Firma Değeri', str(ev)),
            ('52 Hafta Aralığı', str(hafta52)),
            ('Halka Açıklık', str(halka_aciklik)),
            ('Gün. Ort. Hacim', str(gunluk_hacim)),
        ]

    # Piyasa bilgileri başlığı
    p_fb = sag.add_paragraph()
    p_fb.paragraph_format.space_after = Pt(0)
    p_fb.paragraph_format.space_before = Pt(0)
    run_fb = p_fb.add_run('Piyasa Bilgileri')
    _font(run_fb, size=BBB.TABLO_PT, bold=True, color=BBB.BIRINCIL)

    # Key-value tablo — diğer tablolarla aynı tam genişlik
    n_items = len(fiyat_bilgileri)
    fb_tbl = sag.add_table(rows=n_items, cols=2)
    fb_tbl.autofit = False
    fb_tbl.allow_autofit = False
    fb_tbl.columns[0].width = Cm(sag_cm * 0.55)
    fb_tbl.columns[1].width = Cm(sag_cm * 0.45)

    # v4.9: tblLayout fixed + tblW dxa — tüm sağ panel tabloları tutarlı
    _fb_pr = fb_tbl._tbl.tblPr
    for _tag in ('w:tblLayout', 'w:tblW'):
        _ex = _fb_pr.find(qn(_tag))
        if _ex is not None:
            _fb_pr.remove(_ex)
    _fb_layout = OxmlElement('w:tblLayout')
    _fb_layout.set(qn('w:type'), 'fixed')
    _fb_pr.append(_fb_layout)
    _fb_w = OxmlElement('w:tblW')
    _fb_w.set(qn('w:w'), str(int(sag_cm * 567)))
    _fb_w.set(qn('w:type'), 'dxa')
    _fb_pr.append(_fb_w)

    # Tavsiye renk belirle
    _tavsiye_upper = str(tavsiye).upper()
    if _tavsiye_upper in ('AL', 'EKLE', 'BUY', 'ADD'):
        _tavsiye_color = BBB.OLUMLU
    elif _tavsiye_upper in ('SAT', 'AZALT', 'SELL', 'REDUCE'):
        _tavsiye_color = BBB.OLUMSUZ
    else:
        _tavsiye_color = BBB.BIRINCIL

    for i, (etiket, deger) in enumerate(fiyat_bilgileri):
        cell_e = fb_tbl.cell(i, 0)
        cell_v = fb_tbl.cell(i, 1)

        p_e = cell_e.paragraphs[0]
        p_e.paragraph_format.space_after = Pt(0)
        p_e.paragraph_format.space_before = Pt(0)
        run_e = p_e.add_run(etiket)
        _font(run_e, size=BBB.KUCUK_PT, bold=True, color=BBB.GRI)

        p_v = cell_v.paragraphs[0]
        p_v.paragraph_format.space_after = Pt(0)
        p_v.paragraph_format.space_before = Pt(0)
        p_v.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_v = p_v.add_run(str(deger))

        # Tavsiye ve Adil Değer satırlarında vurgu
        if etiket in ('Tavsiye',):
            _font(run_v, size=BBB.TABLO_PT + 1, bold=True, color=_tavsiye_color)
        elif etiket in ('Adil Değer Tahmini', 'Getiri Potansiyeli'):
            _font(run_v, size=BBB.TABLO_PT, bold=True, color=BBB.BIRINCIL)
        else:
            _font(run_v, size=BBB.KUCUK_PT, color=BBB.METIN)

    # v4.8: Zebra kaldırıldı — diğer tablolarla tutarlı stil
    tablo_stil(fb_tbl)
    # v4.9: Cell-içi tablolarda left/right margin sıfır — sağ panel boşluğu düzeltmesi
    _tablo_cell_margin_sifirla(fb_tbl, top=15, bottom=15, left=0, right=0)

    # ─── (2) Finansal Özet Tablosu (tam genişlik, kaynak YOK) ───
    if finansal_ozet:
        p_fo = sag.add_paragraph()
        p_fo.paragraph_format.space_after = Pt(0)
        p_fo.paragraph_format.space_before = Pt(1)
        run_fo = p_fo.add_run('Finansal Özet')
        _font(run_fo, size=BBB.TABLO_PT, bold=True, color=BBB.BIRINCIL)

        tablo_ekle(sag,
                   finansal_ozet.get('basliklar', []),
                   finansal_ozet.get('satirlar', []),
                   kaynak='',
                   kompakt=True, numarala=False,
                   genislik_cm=sag_cm)

    # ─── (3) Çarpanlar Tablosu (tam genişlik, kaynak YOK) ───
    if carpanlar:
        p_ca = sag.add_paragraph()
        p_ca.paragraph_format.space_after = Pt(0)
        p_ca.paragraph_format.space_before = Pt(1)
        run_ca = p_ca.add_run('Değerleme Çarpanları')
        _font(run_ca, size=BBB.TABLO_PT, bold=True, color=BBB.BIRINCIL)

        tablo_ekle(sag,
                   carpanlar.get('basliklar', []),
                   carpanlar.get('satirlar', []),
                   kaynak='',
                   kompakt=True, numarala=False,
                   genislik_cm=sag_cm)

    # ─── (3b) Ortaklık Yapısı Tablosu (opsiyonel) ───
    if ortaklik_yapisi:
        p_oy = sag.add_paragraph()
        p_oy.paragraph_format.space_after = Pt(0)
        p_oy.paragraph_format.space_before = Pt(1)
        run_oy = p_oy.add_run('Ortaklık Yapısı')
        _font(run_oy, size=BBB.TABLO_PT, bold=True, color=BBB.BIRINCIL)

        tablo_ekle(sag,
                   ortaklik_yapisi.get('basliklar', []),
                   ortaklik_yapisi.get('satirlar', []),
                   kaynak='',
                   kompakt=True, numarala=False,
                   genislik_cm=sag_cm)

    # ─── (4) Fiyat Performans Grafiği (EN ALTTA, tam genişlik) ───
    if grafik_yolu and os.path.exists(grafik_yolu):
        p_g = sag.add_paragraph()
        p_g.paragraph_format.space_before = Pt(1)   # v4.12: 3→1pt — sığma optimizasyonu
        p_g.paragraph_format.space_after = Pt(0)
        p_g.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_g = p_g.add_run()
        # v4.12: %90 genişlik — grafiğin yüksekliği orantılı azalır, S.1 sığma sağlar
        run_g.add_picture(grafik_yolu, width=Cm(sag_cm * 0.90))

    # ─── Alt notlar (sağ panel içinde, grafiğin altında) ───
    if alt_notlar:
        for not_metni in alt_notlar:
            p_not = sag.add_paragraph()
            p_not.paragraph_format.space_before = Pt(0)
            p_not.paragraph_format.space_after = Pt(0)
            run_not = p_not.add_run(not_metni)
            _font(run_not, size=6, italic=True, color=BBB.GRI)
    else:
        p_not = sag.add_paragraph()
        p_not.paragraph_format.space_before = Pt(1)
        p_not.paragraph_format.space_after = Pt(0)
        run_not = p_not.add_run(f'*Piyasa verileri {tarih} tarihlidir.')
        _font(run_not, size=6, italic=True, color=BBB.GRI)

    # v4.12: Explicit sayfa sonu kaldırıldı — İçindekiler kendi sayfa_sonu=True ile başlar.
    # Eski _sayfa_sonu(doc) çağrısı boş S.2 yaratıyordu.


# ═══════════════════════════════════════════════════════════
# FİNANSAL TABLO GRİD (Yan Yana)
# ═══════════════════════════════════════════════════════════

def finansal_tablo_grid(doc, sol_basliklar, sol_satirlar, sol_baslik='',
                        sag_basliklar=None, sag_satirlar=None, sag_baslik='',
                        kaynak=''):
    """İki finansal tabloyu yan yana yerleştir.
    
    Örnek: sol=Gelir Tablosu, sag=Nakit Akış Tablosu
    """
    table, sol, sag = iki_sutun_grid(doc, sol_genislik_cm=8.4, sag_genislik_cm=8.4)

    # Sol tablo
    if sol_baslik:
        _paragraf(sol, sol_baslik, size=BBB.TABLO_PT, bold=True, color=BBB.BIRINCIL,
                  space_after=2)
    tablo_ekle(sol, sol_basliklar, sol_satirlar, kompakt=True)

    # Sağ tablo
    if sag_basliklar and sag_satirlar:
        if sag_baslik:
            _paragraf(sag, sag_baslik, size=BBB.TABLO_PT, bold=True, color=BBB.BIRINCIL,
                      space_after=2)
        tablo_ekle(sag, sag_basliklar, sag_satirlar, kompakt=True)

    if kaynak:
        _kaynak(doc, kaynak)

    return table


# ═══════════════════════════════════════════════════════════
# GRAFİK ÖZET SAYFASI (2×3 Grid)
# ═══════════════════════════════════════════════════════════

def grafik_ozet_sayfasi(doc, grafikler, baslik='Temel Grafikler — Görsel Özet'):
    """AK Yatırım tarzı grafik özet sayfası. 2'li grid'de görseller.
    
    grafikler: [{'yol': '...', 'baslik': '...'}, ...] (çift sayıda olması ideal)
    """
    _baslik(doc, baslik, level=1)

    for i in range(0, len(grafikler), 2):
        sol_g = grafikler[i]
        sag_g = grafikler[i + 1] if i + 1 < len(grafikler) else None

        if sag_g:
            cift_grafik(doc,
                        sol_g['yol'], sag_g['yol'],
                        baslik_sol=sol_g.get('baslik', ''),
                        baslik_sag=sag_g.get('baslik', ''))
        else:
            # Tek kalan grafik → tam genişlik
            grafik_ekle(doc, sol_g['yol'], boyut='TAM',
                        baslik=sol_g.get('baslik', ''))

    # NOT: _sayfa_sonu() kaldırıldı — sonraki H1 başlık page_break_before ile yeni sayfa açar


# ─────────────────────────────────────────────────────────────
# TOP-TIER PRIMİTİFLER — T5 Rapor Montajı için yeni bileşenler
# ─────────────────────────────────────────────────────────────

def icgoru_kutusu(doc, metin, tip='bilgi'):
    """Bolum basina key takeaway kutusu. Taranabilirlik artirir.

    tip: 'bilgi' (mavi), 'uyari' (sari), 'olumlu' (yesil), 'risk' (kirmizi)
    """
    renk_map = {
        'bilgi': '#0d579b',
        'uyari': '#FFC000',
        'olumlu': '#329239',
        'risk': '#C00000',
    }
    baslik_map = {
        'bilgi': 'TEMEL BULGU',
        'uyari': 'DİKKAT',
        'olumlu': 'OLUMLU GÖSTERGE',
        'risk': 'ÖNEMLİ RİSK',
    }
    renk = renk_map.get(tip, renk_map['bilgi'])
    baslik_txt = baslik_map.get(tip, baslik_map['bilgi'])

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _hucre_border_modern(cell, top=True, bottom=True, left=True, right=True, color=renk, sz=12)

    # Baslik satiri
    p_baslik = cell.paragraphs[0]
    run = p_baslik.add_run(f'  {baslik_txt}')
    _font(run, size=9, bold=True, color=renk)
    p_baslik.paragraph_format.space_after = Pt(2)  # v4.7 fix: paragraph_format gerekli

    # Icerik
    p_icerik = cell.add_paragraph()
    run = p_icerik.add_run(metin)
    _font(run, size=BBB.METIN_PT, color='#333333')
    p_icerik.paragraph_format.space_after = Pt(3)  # v4.7 fix: paragraph_format gerekli

    # Tablo kenarlarina margin
    _tablo_cell_margin_sifirla(tbl, top=60, bottom=60, left=120, right=120)

    # v4.13: Kutu sonrası spacing — sonraki içerikle arasında nefes alanı
    _spacer = doc.add_paragraph()
    _spacer.paragraph_format.space_before = Pt(0)
    _spacer.paragraph_format.space_after = Pt(4)
    _spacer.paragraph_format.line_spacing = Pt(1)
    _sf = _spacer.add_run()
    _sf.font.size = Pt(1)


def bilgi_kutusu(doc, baslik, icerik_dict, kolonlar=2, vurgu_renk=None):
    """Cerceeveli key metrics kutusu. Kapak sayfasi veya bolum baslari icin.

    icerik_dict: OrderedDict — {'PD': '9 mrd TL', 'EV/EBITDA': '8,0x', ...}
    kolonlar: 2 veya 3
    """
    if vurgu_renk is None:
        vurgu_renk = BBB_RENKLER['birincil']

    items = list(icerik_dict.items())
    satir_sayisi = (len(items) + kolonlar - 1) // kolonlar

    tbl = doc.add_table(rows=satir_sayisi + 1, cols=kolonlar)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Baslik — birlesik hucre
    baslik_cells = tbl.rows[0].cells
    baslik_cells[0].merge(baslik_cells[-1])
    merged = baslik_cells[0]
    _hucre_shading(merged, vurgu_renk)
    p = merged.paragraphs[0]
    run = p.add_run(baslik)
    _font(run, size=10, bold=True, color='#FFFFFF')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Veri satirlari
    for idx, (key, val) in enumerate(items):
        row_idx = idx // kolonlar + 1
        col_idx = idx % kolonlar
        cell = tbl.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        run_key = p.add_run(f'{key}: ')
        _font(run_key, size=9, bold=False, color='#666666')
        run_val = p.add_run(str(val))
        _font(run_val, size=10, bold=True, color='#333333')
        if row_idx % 2 == 0:
            _hucre_shading(cell, BBB.KREM_HEX)

    # Border + orphan koruması
    n_rows = len(tbl.rows)
    for row_idx, row in enumerate(tbl.rows):
        for cell in row.cells:
            _hucre_border_modern(cell, top=True, bottom=True, left=True, right=True,
                                color=vurgu_renk, sz=6)
            # v4.7: orphan koruması — son satır hariç (boş sayfa riski)
            if row_idx < n_rows - 1:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True

    _tablo_cell_margin_sifirla(tbl, top=60, bottom=60, left=100, right=100)
    # v4.7: trailing doc.add_paragraph() kaldırıldı — spacing ile kontrol edilmeli


def skor_karti(doc, skorlar, toplam, seviye_adi, aciklama=None):
    """BBB Kalite Skoru karti. 5 boyut + toplam + seviye.

    skorlar: [{'boyut': 'Moat', 'puan': 4, 'max': 5, 'kantitatif': 'ROIC %38,6', 'kalitatif': 'Genis moat'}, ...]
    """
    max_toplam = sum(s['max'] for s in skorlar)

    tbl = doc.add_table(rows=len(skorlar) + 2, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Baslik satiri
    header_cells = tbl.rows[0].cells
    header_cells[0].merge(header_cells[1])
    p = header_cells[0].paragraphs[0]
    run = p.add_run('BBB KALİTE SKORU')
    _font(run, size=10, bold=True, color='#FFFFFF')
    _hucre_shading(header_cells[0], BBB_RENKLER['birincil'])

    p2 = header_cells[2].paragraphs[0]
    run2 = p2.add_run(f'{toplam}/{max_toplam} — {seviye_adi}')
    _font(run2, size=10, bold=True, color='#FFFFFF')
    _hucre_shading(header_cells[2], BBB_RENKLER['birincil'])
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Boyut satirlari
    for i, skor in enumerate(skorlar):
        row = tbl.rows[i + 1]
        # Boyut adi
        p = row.cells[0].paragraphs[0]
        run = p.add_run(skor['boyut'])
        _font(run, size=9, bold=True, color='#333333')

        # Bar gorseli (Unicode block)
        dolu = '\u2588' * skor['puan']
        bos = '\u2591' * (skor['max'] - skor['puan'])
        p_bar = row.cells[1].paragraphs[0]
        run_bar = p_bar.add_run(f'{dolu}{bos}  {skor["puan"]}/{skor["max"]}')
        _font(run_bar, size=9, color=BBB_RENKLER['birincil'])

        # Kantitatif/kalitatif
        p_detay = row.cells[2].paragraphs[0]
        run_d = p_detay.add_run(f'{skor.get("kantitatif", "")}')
        _font(run_d, size=8, color='#666666', italic=True)

        if i % 2 == 1:
            for cell in row.cells:
                _hucre_shading(cell, BBB.KREM_HEX)

    # Aciklama satiri
    if aciklama:
        son_row = tbl.rows[-1]
        son_row.cells[0].merge(son_row.cells[-1])
        p = son_row.cells[0].paragraphs[0]
        run = p.add_run(aciklama)
        _font(run, size=9, italic=True, color='#666666')

    n_rows = len(tbl.rows)
    for row_idx, row in enumerate(tbl.rows):
        for cell in row.cells:
            _hucre_border_modern(cell, top=True, bottom=True, color=BBB_RENKLER['birincil'], sz=4)
            # v4.7: orphan koruması — son satır hariç
            if row_idx < n_rows - 1:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True

    _tablo_cell_margin_sifirla(tbl, top=40, bottom=40, left=80, right=80)
    # v4.7: trailing doc.add_paragraph() kaldırıldı


def zaman_cizelgesi(doc, olaylar, baslik='Kataliz Takvimi'):
    """Timeline tablo — tarih, olay, beklenen etki.

    olaylar: [{'tarih': 'Q2 2026', 'olay': 'FY2025 sonuclar', 'etki': 'Marj dogrulama', 'onem': 'yuksek'}, ...]
    onem: 'yuksek' (turuncu), 'orta' (gri), 'dusuk' (acik gri)
    """
    onem_renk = {'yuksek': BBB_RENKLER['birincil'], 'orta': '#666666', 'dusuk': '#999999'}
    basliklar_list = ['Tarih', 'Olay / Katalizör', 'Beklenen Etki']

    tbl = doc.add_table(rows=len(olaylar) + 1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Baslik satiri
    for j, b in enumerate(basliklar_list):
        cell = tbl.cell(0, j)
        _hucre_shading(cell, BBB_RENKLER['birincil'])
        p = cell.paragraphs[0]
        run = p.add_run(b)
        _font(run, size=9, bold=True, color='#FFFFFF')

    # Veri satirlari
    for i, olay in enumerate(olaylar):
        renk = onem_renk.get(olay.get('onem', 'orta'), '#666666')
        row = tbl.rows[i + 1]

        # Tarih
        p = row.cells[0].paragraphs[0]
        run = p.add_run(olay['tarih'])
        _font(run, size=9, bold=True, color=renk)

        # Olay
        p = row.cells[1].paragraphs[0]
        run = p.add_run(olay['olay'])
        _font(run, size=9, color='#333333')

        # Etki
        p = row.cells[2].paragraphs[0]
        run = p.add_run(olay['etki'])
        _font(run, size=9, italic=True, color='#666666')

        if i % 2 == 1:
            for cell in row.cells:
                _hucre_shading(cell, BBB.KREM_HEX)

    _tablo_cell_margin_sifirla(tbl, top=40, bottom=40, left=80, right=80)
    # v4.7: trailing doc.add_paragraph() kaldırıldı


def karar_agaci(doc, dallar, baslik='Senaryo Karar Ağacı'):
    """Probability-weighted decision tree tablo formati.
    Maksimum 2 seviye derinlik.

    dallar: [
        {'isim': 'Baz', 'olasilik': 0.50, 'deger': 150, 'alt_dallar': [
            {'isim': 'UK basarili', 'olasilik': 0.60, 'deger': 185},
            {'isim': 'UK notr', 'olasilik': 0.30, 'deger': 155},
        ]},
        {'isim': 'Iyimser', 'olasilik': 0.25, 'deger': 247},
    ]
    """
    # Satirlari duzlestir
    rows = []
    for dal in dallar:
        pct = f'%{dal["olasilik"]*100:.0f}'
        rows.append((dal['isim'], pct, f'{dal["deger"]:.0f} TL', 0))
        for alt in dal.get('alt_dallar', []):
            alt_pct = f'%{alt["olasilik"]*100:.0f}'
            rows.append((alt['isim'], alt_pct, f'{alt["deger"]:.0f} TL', 1))

    # Olasilik agirlikli deger hesapla
    agirlikli = sum(d['olasilik'] * d['deger'] for d in dallar)

    tbl = doc.add_table(rows=len(rows) + 2, cols=3)  # +1 baslik +1 sonuc
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Baslik
    header = tbl.rows[0].cells
    header[0].merge(header[-1])
    _hucre_shading(header[0], BBB_RENKLER['birincil'])
    p = header[0].paragraphs[0]
    run = p.add_run(baslik)
    _font(run, size=10, bold=True, color='#FFFFFF')

    # Dallar
    for i, (isim, pct, deger, seviye) in enumerate(rows):
        row = tbl.rows[i + 1]
        indent = '    ' if seviye == 1 else ''
        prefix = '├ ' if seviye == 1 else ''

        p = row.cells[0].paragraphs[0]
        run = p.add_run(f'{indent}{prefix}{isim}')
        _font(run, size=9, bold=(seviye==0), color='#333333')

        p = row.cells[1].paragraphs[0]
        run = p.add_run(f'({pct})')
        _font(run, size=9, color='#666666')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = row.cells[2].paragraphs[0]
        run = p.add_run(f'→ {deger}')
        _font(run, size=9, bold=True, color=BBB_RENKLER['birincil'] if seviye==0 else '#666666')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Sonuc satiri
    son = tbl.rows[-1].cells
    son[0].merge(son[1])
    _hucre_shading(son[0], BBB.ACIK_HEX)
    p = son[0].paragraphs[0]
    run = p.add_run('Olasılık Ağırlıklı Değer:')
    _font(run, size=9, bold=True, color='#333333')

    _hucre_shading(son[2], BBB.ACIK_HEX)
    p = son[2].paragraphs[0]
    run = p.add_run(f'{agirlikli:.0f} TL')
    _font(run, size=10, bold=True, color=BBB_RENKLER['birincil'])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    _tablo_cell_margin_sifirla(tbl, top=40, bottom=40, left=80, right=80)
    # v4.7: trailing doc.add_paragraph() kaldırıldı


def ozet_finansal(doc, yillar, metrikler, baslik='Özet Finansal Tablo'):
    """Kompakt cok-yillik P&L / Bilanco / CF tablosu.

    yillar: ['FY2023', 'FY2024', 'FY2025', 'FY2026T', 'FY2027T']
    metrikler: OrderedDict — {'Hasilat (mn TL)': [val, ...], 'EBITDA (mn TL)': [...], ...}
    Tahmin yillari T soneki ile isaretlenir ve farkli stil alir.
    """
    basliklar_list = [''] + yillar
    satirlar = []
    for metrik_adi, degerler in metrikler.items():
        satir = [metrik_adi]
        for d in degerler:
            if isinstance(d, (int, float)):
                satir.append(f'{d:,.0f}')
            else:
                satir.append(str(d))
        satirlar.append(satir)

    tbl = doc.add_table(rows=len(satirlar) + 1, cols=len(basliklar_list))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Baslik satiri
    for j, b in enumerate(basliklar_list):
        cell = tbl.cell(0, j)
        _hucre_shading(cell, BBB_RENKLER['birincil'])
        p = cell.paragraphs[0]
        run = p.add_run(b)
        _font(run, size=9, bold=True, color='#FFFFFF')
        if j > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Veri satirlari
    for i, satir in enumerate(satirlar):
        row = tbl.rows[i + 1]
        for j, val in enumerate(satir):
            cell = row.cells[j]
            p = cell.paragraphs[0]

            # Tahmin yili kontrolu
            is_tahmin = j > 0 and yillar[j-1].endswith('T')

            run = p.add_run(val)
            if j == 0:
                _font(run, size=9, bold=True, color='#333333')
            elif is_tahmin:
                _font(run, size=9, italic=True, color=BBB_RENKLER['birincil'])
            else:
                _font(run, size=9, color='#333333')

            if j > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if i % 2 == 1:
                _hucre_shading(cell, BBB.KREM_HEX)

    _tablo_cell_margin_sifirla(tbl, top=30, bottom=30, left=60, right=60)
    # v4.7: trailing doc.add_paragraph() kaldırıldı


# ═══════════════════════════════════════════════════════════
# YASAL UYARI
# ═══════════════════════════════════════════════════════════

def yasal_uyari(doc):
    """Son sayfa: yasal uyarı. v4.14: SPK uyumlu metin."""
    # v4.14: sayfa_sonu geri eklendi — Yasal Uyarı her zaman yeni sayfada başlamalı
    _baslik(doc, 'Yasal Uyarı', level=1, sayfa_sonu=True)
    metinler = [
        "Burada yer alan bilgiler genel bilgilendirme amacı ile hazırlanmıştır. Yatırım Danışmanlığı "
        "hizmeti; aracı kurumlar, portföy yönetim şirketleri, mevduat kabul etmeyen bankalar ile müşteri "
        "arasında imzalanacak Yatırım Danışmanlığı sözleşmesi çerçevesinde sunulmaktadır. Burada yer alan "
        "yorum ve tavsiyeler yorum ve tavsiyede bulunanların kişisel görüşlerine dayanmakta olup, herhangi "
        "bir yatırım aracının alım-satım önerisi ya da getiri vaadi olarak yorumlanmamalıdır. Bu görüşler "
        "mali durumunuz ile risk ve getiri tercihlerinize uygun olmayabilir. Bu nedenle, sadece burada yer "
        "alan bilgilere dayanarak yatırım kararı verilmesi beklentilerinize uygun sonuçlar doğurmayabilir. "
        "Hiçbir şekilde yönlendirici nitelikte olmayan bu içerik, genel anlamda bilgi vermeyi amaçlamakta "
        "olup; bu içeriğin, kişilerin yatırımcıların alım satım kararlarını destekleyebilecek yeterli "
        "bilgileri kapsamayabileceği dikkate alınmalıdır. Bu raporda yer alan çeşitli bilgi ve görüşlere "
        "dayanılarak yapılacak ileriye dönük yatırımlar ve ticari işlemlerin sonuçlarından ya da ortaya "
        "çıkabilecek zararlardan içerik sahibi sorumlu tutulamaz.",

        "Burada yer alan fiyatlar, veriler ve bilgilerin tam ve doğru olduğu garanti edilemez; içerik, "
        "haber verilmeksizin değiştirilebilir. Tüm veriler, borsadabibasina.com tarafından güvenilir "
        "olduğuna inanılan kaynaklardan alınmış ve kaynak belirtilmeyen veriler hesaplanarak paylaşılmıştır. "
        "Hesaplama hatası veya kaynakların kullanılması nedeni ile ortaya çıkabilecek hatalardan içerik "
        "üreticisi sorumlu değildir. Bu raporda yer alan içeriğin tamamı veya herhangi bir kısmı yalnızca "
        "sahipleri tarafından veya sahiplerinin yazılı izni ile kullanılabilir. İçerik üzerinde izinsiz "
        "değişiklik yapmak, kopyalamak, kiralamak, ödünç vermek, iletmek ve yayınlamak yasaktır. Bu "
        "rapordan alınan içerik herhangi bir ticari amaçla kullanılamaz.",
    ]
    for m in metinler:
        _paragraf(doc, m, size=BBB.KAYNAK_PT, color=BBB.GRI, space_after=8)


# ═══════════════════════════════════════════════════════════
# DOCX BAŞLATMA (Ortak Setup)
# ═══════════════════════════════════════════════════════════

def _doc_olustur():
    """Yeni DOCX oluştur, varsayılan stiller ayarla. Sıcak Premium tasarım dili."""
    sayac_sifirla()  # Grafik/Tablo numaralama sıfırla
    doc = Document()

    # ── Sıcak Premium: Sayfa arka plan rengi (bej) ──
    # Katman 1: w:background — explicit renk (themeColor YOK — override'ı önler)
    background = OxmlElement('w:background')
    background.set(qn('w:color'), BBB.SAYFA_BG_HEX)
    # VML fill — Word'ün tüm sürümlerinde (Mac/Win) güvenilir arka plan
    vml_bg = parse_xml(
        f'<v:background xmlns:v="urn:schemas-microsoft-com:vml">'
        f'<v:fill type="solid" color="#{BBB.SAYFA_BG_HEX}"/>'
        f'</v:background>'
    )
    background.append(vml_bg)
    doc.element.insert(0, background)

    # Katman 2: displayBackgroundColors — Word'e arka planı göstermesini söyler
    try:
        settings = doc.settings.element
        # Varsa kaldır, tekrar ekle (idempotent)
        existing = settings.find(qn('w:displayBackgroundColors'))
        if existing is not None:
            settings.remove(existing)
        display_bg = OxmlElement('w:displayBackgroundColors')
        settings.insert(0, display_bg)  # İlk eleman olarak ekle — öncelik
    except Exception:
        pass

    # Katman 3: Tema rengi override — background1 (lt1) rengini bej yap
    # python-docx tema part'ını generic Part olarak tutar, blob üzerinden XML manipülasyonu gerekir
    try:
        from lxml import etree as _etree
        theme_part = doc.part.part_related_by(
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme')
        theme_xml = _etree.fromstring(theme_part.blob)
        ns = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
        lt1 = theme_xml.find('.//a:themeElements/a:clrScheme/a:lt1', ns)
        if lt1 is not None:
            for child in list(lt1):
                lt1.remove(child)
            srgb = _etree.SubElement(lt1, f'{{{ns["a"]}}}srgbClr')
            srgb.set('val', BBB.SAYFA_BG_HEX)
            theme_part._blob = _etree.tostring(theme_xml, xml_declaration=True,
                                                encoding='UTF-8', standalone=True)
    except Exception:
        pass  # Tema erişimi başarısızsa Katman 1+2 yeterli

    # Varsayılan font + line spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = BBB.FONT
    font.size = Pt(BBB.METIN_PT)
    font.color.rgb = BBB.METIN
    style.paragraph_format.line_spacing = Pt(10.5)  # ~1.17× for 9pt — top-tier sıkı yoğunluk
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    style.paragraph_format.space_after = Pt(2)  # v4.7: 3→2pt — sayfa doluluk optimizasyonu
    style.paragraph_format.widow_control = True  # Yetim/dul satır koruması — tek satır taşmayı önler

    # Sayfa boyutu ve kenar boşlukları (A4 — asimetrik)
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(BBB.UST_KENAR_CM)
        sec.bottom_margin = Cm(BBB.ALT_KENAR_CM)
        sec.left_margin = Cm(BBB.SOL_KENAR_CM)
        sec.right_margin = Cm(BBB.SAG_KENAR_CM)

    return doc


# ═══════════════════════════════════════════════════════════
# ŞABLON OLUŞTURMA
# ═══════════════════════════════════════════════════════════

def sablon_olustur(rapor_turu, ticker, cikti, sirket_adi='', tarih=''):
    """Profesyonel DOCX şablon oluştur."""
    doc = _doc_olustur()

    # Header/Footer
    header_text = f'Şirket Raporu — {sirket_adi or ticker}'
    header_footer_ekle(doc, header_text)

    # Sayfa 1: Kapak
    kapak_sayfasi(doc, ticker, sirket_adi=sirket_adi, tarih=tarih,
                  rapor_turu={
                      'initiation': 'KAPSAM BAŞLATMA',
                      'earnings': 'ÇEYREKLİK GÜNCELLEME',
                      'sektor': 'SEKTÖR RAPORU',
                  }.get(rapor_turu, 'ANALİZ'))

    # İçindekiler
    icindekiler_ekle(doc)
    _sayfa_sonu(doc)

    # Bölüm başlıkları (rapor türüne göre)
    bolumler = {
        'initiation': [
            'Yatırım Tezi ve Riskler',
            'Şirket Profili',
            'Yönetim Analizi',
            'Sektör Analizi',
            'Moat Analizi',
            'Finansal Analiz',
            'Gelir Modeli ve Projeksiyonlar',
            'Projeksiyon Varsayımları',
            'Senaryo Analizi',
            'Değerleme Analizi',
            'Peer Karşılaştırma',
            'Thesis Scorecard',
        ],
        'earnings': [
            'Çeyreklik Sonuç Özeti',
            'Beat/Miss Analizi',
            'Bu Çeyrekte Ne Değişti?',
            'Tahmin Revizyonu',
            'Teze Etkisi',
            'Değerleme Güncelleme',
        ],
        'sektor': [
            'Sektör Genel Bakış',
            'Pazar Büyüklüğü ve Yapı',
            'Porter Analizi',
            'Düzenleyici Ortam',
            'Rekabet Haritası',
            'Sektör Değerleme',
            'Yatırım İmplikasyonları',
        ],
    }.get(rapor_turu, ['Analiz'])

    for bolum in bolumler:
        _baslik(doc, bolum, level=1)
        _paragraf(doc, f'[{bolum} içeriği buraya yazılacak]', italic=True, color=BBB.GRI)
        _bos_satir(doc)

    # Yasal uyarı
    yasal_uyari(doc)

    # Kaydet
    Path(cikti).parent.mkdir(parents=True, exist_ok=True)
    doc.save(cikti)
    print(f'✅ Şablon oluşturuldu → {cikti}')
    return doc


# ═══════════════════════════════════════════════════════════
# MARKDOWN → DOCX (Gelişmiş Parser)
# ═══════════════════════════════════════════════════════════

def _parse_inline(paragraph, text):
    """Inline markdown: **bold**, *italic*, ***bold italic*** parse et."""
    # Pattern: ***text*** → bold italic
    # Pattern: **text** → bold
    # Pattern: *text* → italic
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))')
    for match in pattern.finditer(text):
        if match.group(2):  # ***bold italic***
            run = paragraph.add_run(match.group(2))
            _font(run, bold=True, italic=True, color=BBB.METIN)
        elif match.group(3):  # **bold**
            run = paragraph.add_run(match.group(3))
            _font(run, bold=True, color=BBB.METIN)
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            _font(run, italic=True, color=BBB.METIN)
        elif match.group(5):  # normal text
            run = paragraph.add_run(match.group(5))
            _font(run, color=BBB.METIN)


def markdown_to_docx(md_path, ticker, cikti, grafikler_klasor=None,
                     sirket_adi='', tarih=''):
    """Markdown'ı profesyonel DOCX'e dönüştür."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = _doc_olustur()
    header_footer_ekle(doc, f'Şirket Raporu — {sirket_adi or ticker}')

    lines = content.split('\n')
    in_table = False
    table_rows = []
    in_code = False
    in_list = False

    for line in lines:
        stripped = line.strip()

        # Code block
        if stripped.startswith('```'):
            in_code = not in_code
            continue
        if in_code:
            _paragraf(doc, line, size=BBB.TABLO_PT, color=BBB.GRI)
            continue

        # Tablo biriktirme
        if stripped.startswith('|') and '|' in stripped[1:]:
            if re.match(r'^\|[\s\-:]+\|', stripped):
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            continue
        elif in_table:
            if table_rows:
                basliklar = table_rows[0]
                satirlar = table_rows[1:]
                tablo_ekle(doc, basliklar, satirlar)
            in_table = False
            table_rows = []

        # Boş satır
        if not stripped:
            in_list = False
            continue

        # Headings
        if stripped.startswith('# '):
            _baslik(doc, stripped[2:], level=1)
        elif stripped.startswith('## '):
            _baslik(doc, stripped[3:], level=2)
        elif stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=3)
            for run in h.runs:
                _font(run, size=BBB.METIN_PT, bold=True, color=BBB.BIRINCIL)
        # Bullet list
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _parse_inline(p, stripped[2:])
            in_list = True
        # Numbered list
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            _parse_inline(p, re.sub(r'^\d+\.\s', '', stripped))
            in_list = True
        # Kaynak notu
        elif stripped.lower().startswith('kaynak:'):
            _kaynak(doc, stripped)
        # Normal paragraf
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _parse_inline(p, stripped)

    # Son tablo
    if in_table and table_rows:
        basliklar = table_rows[0]
        satirlar = table_rows[1:]
        tablo_ekle(doc, basliklar, satirlar)

    # Grafikleri göm (klasörden)
    if grafikler_klasor and os.path.isdir(grafikler_klasor):
        grafik_dosyalari = sorted([
            f for f in os.listdir(grafikler_klasor)
            if f.lower().endswith(('.png', '.jpg', '.jpeg'))
        ])
        if grafik_dosyalari:
            _sayfa_sonu(doc)
            _baslik(doc, 'Grafikler', level=1)

            # 2'li grid'de göm
            for i in range(0, len(grafik_dosyalari), 2):
                sol_f = grafik_dosyalari[i]
                sag_f = grafik_dosyalari[i + 1] if i + 1 < len(grafik_dosyalari) else None

                sol_yol = os.path.join(grafikler_klasor, sol_f)
                sol_baslik = sol_f.replace('.png', '').replace('.jpg', '').replace('_', ' ').title()

                if sag_f:
                    sag_yol = os.path.join(grafikler_klasor, sag_f)
                    sag_baslik = sag_f.replace('.png', '').replace('.jpg', '').replace('_', ' ').title()
                    cift_grafik(doc, sol_yol, sag_yol,
                                baslik_sol=sol_baslik, baslik_sag=sag_baslik)
                else:
                    grafik_ekle(doc, sol_yol, boyut='TAM', baslik=sol_baslik)
                _bos_satir(doc)

    # Yasal uyarı
    yasal_uyari(doc)

    # Kaydet
    Path(cikti).parent.mkdir(parents=True, exist_ok=True)
    doc.save(cikti)

    p_count = len(doc.paragraphs)
    t_count = len(doc.tables)
    print(f'✅ DOCX oluşturuldu → {cikti}')
    print(f'   Paragraf: {p_count} | Tablo: {t_count}')
    return doc


# ═══════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description='BBB DOCX Rapor Üretici v3.0 — Kurum seviyesi profesyonel layout',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Örnekler:
  # AK Yatırım tarzı initiation şablonu
  python rapor-uret.py --sablon initiation --ticker TBORG --sirket "Türk Tuborg" --cikti rapor.docx

  # Markdown → DOCX
  python rapor-uret.py --markdown analysis.md --ticker TBORG --cikti rapor.docx

  # Grafik embed ile
  python rapor-uret.py --markdown analysis.md --ticker TBORG --grafikler charts/ --cikti rapor.docx

Şablon türleri: initiation, earnings, sektor

Layout özellikleri (v2.0):
  - AK Yatırım tarzı dual-column sayfa 1
  - 2'li grafik grid (yan yana chart yerleşimi)
  - Finansal tablo grid (yan yana tablolar)
  - Header/footer (şirket adı + sayfa no)
  - Gerçek Word İçindekiler (TOC field)
  - Grafik boyut kontrolü: TAM/YARIM/GRID/KUCUK
        """
    )
    parser.add_argument('--sablon', choices=['initiation', 'earnings', 'sektor'],
                        help='Boş şablon türü')
    parser.add_argument('--markdown', help='Markdown dosya yolu')
    parser.add_argument('--ticker', required=True, help='Hisse kodu')
    parser.add_argument('--sirket', default='', help='Şirket adı')
    parser.add_argument('--tarih', default='', help='Rapor tarihi')
    parser.add_argument('--grafikler', help='Grafik klasörü (PNG dosyaları)')
    parser.add_argument('--cikti', required=True, help='Çıktı DOCX dosya yolu')

    args = parser.parse_args()

    if args.sablon:
        sablon_olustur(args.sablon, args.ticker, args.cikti,
                       sirket_adi=args.sirket, tarih=args.tarih)
    elif args.markdown:
        markdown_to_docx(args.markdown, args.ticker, args.cikti,
                         grafikler_klasor=args.grafikler,
                         sirket_adi=args.sirket, tarih=args.tarih)
    else:
        parser.print_help()


if __name__ == '__main__':
    main()
